#!/usr/bin/env python3
"""
Dossiê Estratégico de Defesa — Tribunal do Júri
Cosme Miguel Gomes de Souza e Marcélio Miguel de Souza
Processo: 8006374-21.2024.8.05.0039

Padrão Defender v2 — Paleta Emerald (Júri)
Gerado automaticamente pelo OMBUDS / Claude Code
"""

import os
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, Twips, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import numpy as np

# ============================================================
# PALETA EMERALD (JÚRI)
# ============================================================
DARK_HEADER = "1A5C36"
ACCENT = "2D8B57"
SUBTLE_BG = "F0FFF4"
SOFT_HIGHLIGHT = "D1FAE5"
TITLE_TEXT = "064E3B"
MUTED = "4A7A5E"
WHITE = "FFFFFF"
LIGHT_BORDER = "E5E5E5"
WARNING_BG = "FFFBEB"
ERROR_BG = "FEE2E2"
SUCCESS_BG = "D1FAE5"
PENDING_BG = "FEF9C3"
WARNING_BORDER = "C8A84E"
ERROR_BORDER = "DC2626"

# ============================================================
# PATHS
# ============================================================
LOGO_PATH = "/Users/rodrigorochameire/Projetos/Defender/.claude/skills-cowork/juri/assets/dpe_logo.png"
OUTPUT_DIR = "/Users/rodrigorochameire/Meu Drive/1 - Defensoria 9ª DP/Processos - Júri/Cosme Miguel Gomes de Souza e Marcelio Miguel de Souza"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "DOSSIE_ESTRATEGICO_COSME_MARCELIO.docx")

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_left_border(cell, color, width='18'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), width)
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color)
    borders.append(left)
    tcPr.append(borders)

def set_cell_borders(cell, color=LIGHT_BORDER):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        borders.append(border)
    tcPr.append(borders)

def set_table_borders(table, color=LIGHT_BORDER):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, color)

def add_header_row(table, texts, bg=DARK_HEADER, font_color=WHITE):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(font_color)
        run.font.name = 'Verdana'

def add_data_row(table, row_idx, texts, bg=None):
    row = table.rows[row_idx]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        if bg:
            set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Verdana'

def add_section_header(doc, text, level=1):
    """Add section header with left accent border"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_left_border(cell, ACCENT, '24')
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(14)
    else:
        run.bold = True
        run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(TITLE_TEXT)
    run.font.name = 'Verdana'
    doc.add_paragraph()

def add_subsection(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TITLE_TEXT)
    run.font.name = 'Verdana'

def add_body(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Verdana'
    run.italic = italic
    run.bold = bold
    p.paragraph_format.line_spacing = 1.15
    return p

def add_citation(doc, text, source=""):
    """Add a citation in italics with source"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f'"{text}"')
    run.font.size = Pt(10)
    run.font.name = 'Verdana'
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    if source:
        run2 = p.add_run(f' — {source}')
        run2.font.size = Pt(9)
        run2.font.name = 'Verdana'
        run2.font.color.rgb = RGBColor.from_string(MUTED)
    return p

def add_alert_box(doc, text, bg=WARNING_BG, border=WARNING_BORDER):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_left_border(cell, border, '24')
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Verdana'
    doc.add_paragraph()

def add_separator(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), ACCENT)
    borders.append(bottom)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

def add_card(doc, title, content_lines, border_color=ACCENT):
    """Add a card with left border"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_left_border(cell, border_color, '18')
    set_cell_bg(cell, WHITE)
    # Title
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TITLE_TEXT)
    run.font.name = 'Verdana'
    # Content
    for line in content_lines:
        p2 = cell.add_paragraph()
        if line.startswith('_') and line.endswith('_'):
            run2 = p2.add_run(line[1:-1])
            run2.italic = True
            run2.font.color.rgb = RGBColor.from_string(MUTED)
        else:
            run2 = p2.add_run(line)
        run2.font.size = Pt(10)
        run2.font.name = 'Verdana'
    doc.add_paragraph()

def fade_logo(input_path, output_path, opacity=0.60):
    try:
        img = Image.open(input_path).convert("RGBA")
        arr = np.array(img, dtype=np.float64)
        white = np.full_like(arr[:,:,:3], 255.0)
        arr[:,:,:3] = arr[:,:,:3] * opacity + white * (1 - opacity)
        arr[:,:,3] = 255
        Image.fromarray(arr.astype(np.uint8)).convert("RGB").save(output_path)
        return True
    except Exception as e:
        print(f"Aviso: Não foi possível processar logo: {e}")
        return False

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Defensoria Pública do Estado da Bahia")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("7ª Regional da DPE – Camaçari – Bahia.")
    run2.font.name = 'Arial Narrow'
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor.from_string(MUTED)


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def generate_dossie():
    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    style.font.name = 'Verdana'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Margins
    section = doc.sections[0]
    section.top_margin = Twips(2552)
    section.bottom_margin = Twips(1134)
    section.left_margin = Twips(1418)
    section.right_margin = Twips(1134)

    # Header with faded logo
    faded_path = "/tmp/dpe_logo_faded.png"
    if fade_logo(LOGO_PATH, faded_path):
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run()
        run.add_picture(faded_path, width=Inches(1.777))

    # Footer
    add_footer(doc)

    # ========================================================
    # CAPA + DASHBOARD
    # ========================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOSSIÊ ESTRATÉGICO DE DEFESA")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(TITLE_TEXT)
    run.font.name = 'Verdana'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("TRIBUNAL DO JÚRI")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor.from_string(ACCENT)
    run2.font.name = 'Verdana'

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Cosme Miguel Gomes de Souza\ne\nMarcélio Miguel de Souza")
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor.from_string(TITLE_TEXT)
    run3.font.name = 'Verdana'

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Processo nº 8006374-21.2024.8.05.0039\nGerado em {datetime.now().strftime('%d/%m/%Y')}")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor.from_string(MUTED)
    run4.font.name = 'Verdana'

    doc.add_paragraph()

    # KPI Dashboard — 4 cards
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ("4+ anos", "Tempo preso\nCosme"),
        ("4+ anos", "Tempo preso\nMarcélio"),
        ("6", "Testemunhas\nouvidas"),
        ("Legítima\nDefesa", "Tese\nprincipal")
    ]
    for i, (value, label) in enumerate(kpis):
        # Value
        cell_v = kpi_table.rows[0].cells[i]
        set_cell_bg(cell_v, SUBTLE_BG)
        set_cell_borders(cell_v, ACCENT)
        pv = cell_v.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(value)
        rv.bold = True
        rv.font.size = Pt(24)
        rv.font.color.rgb = RGBColor.from_string(DARK_HEADER)
        rv.font.name = 'Verdana'
        # Label
        cell_l = kpi_table.rows[1].cells[i]
        set_cell_bg(cell_l, WHITE)
        set_cell_borders(cell_l, ACCENT)
        pl = cell_l.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(label)
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor.from_string(MUTED)
        rl.font.name = 'Verdana'

    doc.add_paragraph()

    # Info strip
    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Processo", "8006374-21.2024.8.05.0039"),
        ("Vara", "Júri e Exec. Penais — Camaçari"),
        ("Status", "Aguardando Plenário")
    ]
    for i, (label, value) in enumerate(info_data):
        cell = info_table.rows[0].cells[i]
        set_cell_bg(cell, SOFT_HIGHLIGHT)
        set_cell_borders(cell, ACCENT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.name = 'Verdana'
        r2 = p.add_run(value)
        r2.font.size = Pt(9)
        r2.font.name = 'Verdana'

    doc.add_paragraph()

    # ÍNDICE
    add_subsection(doc, "SUMÁRIO")
    indice = [
        "PARTE I — VISÃO GERAL",
        "PARTE II — O CASO",
        "PARTE III — PROVA",
        "PARTE IV — ESTRATÉGIA",
        "PARTE V — PLENÁRIO",
        "PARTE VI — CENÁRIOS",
        "PARTE VII — PROVIDÊNCIAS"
    ]
    for item in indice:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Verdana'
        run.font.color.rgb = RGBColor.from_string(ACCENT)

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE I — VISÃO GERAL
    # ========================================================
    add_section_header(doc, "PARTE I — VISÃO GERAL")

    # 1.1 Resumo Executivo
    add_subsection(doc, "1.1 Resumo Executivo")

    add_body(doc, "Cosme Miguel Gomes de Souza e Marcélio Miguel de Souza são acusados da prática de homicídio (art. 121 do Código Penal) contra Cláudio de Jesus Silva, 39 anos, trabalhador da construção civil, fato ocorrido em 28 de julho de 2018, no Condomínio Pardais III, Bloco 82, Jardim Limoeiro, Camaçari-BA. A denúncia sustenta que os defendidos, irmãos da companheira da suposta vítima, teriam desferido golpes de faca que resultaram no óbito. Ambos os defendidos encontram-se presos há mais de 4 anos — Cosme desde 29/12/2021 e Marcélio desde 16/01/2022.")

    add_body(doc, "O acervo probatório apresenta fragilidades significativas para a acusação. A principal testemunha presencial — Elisangela Maria Gomes de Souza, irmã dos defendidos e companheira da suposta vítima — apresenta contradições relevantes entre seus depoimentos. Na versão policial de 2017, Elisangela declarou que Cosme apenas tomou a faca da suposta vítima, sem desferir golpes. Já em 2023, perante outro delegado, afirmou que ambos os irmãos esfaquearam Cláudio. As demais testemunhas são vizinhos que não presenciaram os fatos diretamente, e os policiais civis não se recordam da ocorrência.")

    add_body(doc, "A tese defensiva principal é a LEGÍTIMA DEFESA PRÓPRIA E DE TERCEIRO. As provas indicam que Cláudio — que possuía histórico de violência doméstica, ameaças e agressões contra a família — foi o agressor inicial: pegou uma faca (peixeira) na cozinha, trancou-se no quarto aguardando Marcélio, e atacou primeiro. Marcélio foi cortado no braço antes de reagir. Cosme foi chamado para separar a briga e também foi atingido. Subsidiariamente, sustenta-se excesso culposo na legítima defesa e homicídio privilegiado por violenta emoção.")

    # 1.2 Contatos
    add_subsection(doc, "1.2 Contatos")
    ct = doc.add_table(rows=6, cols=3)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(ct, ["Pessoa/Órgão", "Função", "Contato"])
    contacts = [
        ("Cosme Miguel Gomes de Souza", "Defendido 1 (preso)", "Conj. Penal de Juazeiro"),
        ("Marcélio Miguel de Souza", "Defendido 2 (preso)", "Colônia Penal Simões Filho"),
        ("Elisangela Maria Gomes de Souza", "Irmã / informante", "[VERIFICAR TELEFONE]"),
        ("Vara do Júri — Camaçari", "Juízo", "[VERIFICAR TELEFONE]"),
        ("Promotoria de Justiça", "MP — Dr. Gilberto", "[VERIFICAR TELEFONE]"),
    ]
    for i, (pessoa, funcao, contato) in enumerate(contacts):
        add_data_row(ct, i+1, [pessoa, funcao, contato], SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(ct)
    doc.add_paragraph()

    # 1.3 Prazos em aberto
    add_subsection(doc, "1.3 Prazos em Aberto")
    pt = doc.add_table(rows=4, cols=4)
    add_header_row(pt, ["Data", "Status", "Descrição", "Observação"])
    prazos = [
        ("[A DEFINIR]", "⏰", "Plenário do Júri", "Aguardando pauta"),
        ("[VERIFICAR]", "🔴", "HC — excesso de prazo (4+ anos preso)", "Urgente — avaliar impetração"),
        ("[VERIFICAR]", "⏰", "Requisição para plenário", "Verificar se já requisitados"),
    ]
    for i, (data, status, desc, obs) in enumerate(prazos):
        add_data_row(pt, i+1, [data, status, desc, obs])
    set_table_borders(pt)
    doc.add_paragraph()

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE II — O CASO
    # ========================================================
    add_section_header(doc, "PARTE II — O CASO")

    # 2.1 Partes e qualificação
    add_subsection(doc, "2.1 Partes e Qualificação")

    add_card(doc, "DEFENDIDO 1 — Cosme Miguel Gomes de Souza", [
        "Idade: 41 anos | Estado civil: Solteiro",
        "Filhos: 9 (1 menor de 11 anos)",
        "Profissão: Trabalhador braçal (coletor de lixo, obras)",
        "Residência anterior: Miradima/Pernambuco, Bairro Projetada",
        "Situação: PRESO desde 29/12/2021 — Conjunto Penal de Juazeiro",
        "Tempo preso: 4 anos e 3 meses (aprox.)",
        "Antecedentes: [VERIFICAR]",
    ], ACCENT)

    add_card(doc, "DEFENDIDO 2 — Marcélio Miguel de Souza", [
        "Idade: 24 anos | Estado civil: Solteiro",
        "Filhos: 1 (14 anos, com anemia — sem tratamento por estar preso)",
        "Profissão: Trabalhador braçal (frigorífico, construção, roça)",
        "Residência anterior: Salvador (com irmã) / Feira de Santana",
        "Situação: PRESO desde 16/01/2022 — Colônia Penal de Simões Filho",
        "Tempo preso: 4 anos e 2 meses (aprox.)",
        "Antecedentes: [VERIFICAR]",
    ], ACCENT)

    add_card(doc, "SUPOSTA VÍTIMA — Cláudio de Jesus Silva (falecido)", [
        "Idade ao óbito: 39 anos | Apelido: 'Negão'",
        "Profissão: Trabalhador da construção civil / caseiro em roça",
        "Relação com defendidos: Companheiro de Elisangela (irmã dos defendidos)",
        "Filhos com Elisangela: 3 (à época, menores de idade)",
        "Histórico: Violência doméstica reiterada, ameaças de morte à companheira e cunhados",
        "_'Ele me batia todos os dias sem eu fazer nada, me traia, botava ele em casa' — Elisangela, audiência judicial_",
        "_'Ele já vinha ameaçando a gente já havia mais de cinco anos' — Cosme, interrogatório_",
        "Antecedentes: Ocorrências por VD contra companheiras (confirmado por Elisangela)",
    ], ERROR_BORDER)

    # 2.2 Fato imputado
    add_subsection(doc, "2.2 Fato Imputado")
    add_body(doc, "Segundo a denúncia do Ministério Público, no dia 28 de julho de 2018, por volta das 18h30, no interior do apartamento situado no Condomínio Pardais III, Bloco 82, Jardim Limoeiro, Camaçari-BA, os defendidos Cosme e Marcélio teriam, em concurso de pessoas, desferido golpes de arma branca (faca) contra Cláudio de Jesus Silva, causando-lhe lesões que resultaram em óbito.")
    add_body(doc, "Tipificação: Art. 121, caput, c/c art. 29, ambos do Código Penal.")

    add_alert_box(doc, "⚠ PONTO CRÍTICO: A denúncia não reconhece que Cláudio foi o agressor inicial. As provas nos autos demonstram que Cláudio pegou a faca primeiro, trancou-se esperando Marcélio, e atacou primeiro. A acusação ignora o histórico de 5+ anos de violência e ameaças.", ERROR_BG, ERROR_BORDER)

    # 2.3 Cronologia processual
    add_subsection(doc, "2.3 Cronologia Processual")
    crono = doc.add_table(rows=11, cols=3)
    add_header_row(crono, ["Data", "Evento", "Status"])
    eventos = [
        ("28/07/2018", "🔴 Fato: morte de Cláudio no Cond. Pardais III", ""),
        ("28/07/2018", "⭐ SILC comparece ao local — sangue por todo lado", ""),
        ("01/11/2017*", "Depoimento Elisangela (Deleg. Geovane) — Cosme só tomou a faca", ""),
        ("2018", "Primeiro depoimento policial — Elisangela (inquérito)", ""),
        ("2023", "Reintimação Elisangela (Deleg. Almir) — muda versão: ambos esfaquearam", "🟡"),
        ("29/12/2021", "🔴 Prisão de Cosme — Conjunto Penal de Juazeiro", ""),
        ("16/01/2022", "🔴 Prisão de Marcélio — Colônia Penal Simões Filho", ""),
        ("2024", "Distribuição do processo — Vara do Júri Camaçari", ""),
        ("2025", "⭐ Audiência de instrução — oitiva de testemunhas e interrogatórios", ""),
        ("[A DEFINIR]", "🟢 Plenário do Júri", "PENDENTE"),
    ]
    for i, (data, evento, status) in enumerate(eventos):
        add_data_row(crono, i+1, [data, evento, status], SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(crono)
    doc.add_paragraph()

    add_alert_box(doc, "⚠ NOTA: Há referência a depoimento datado de 01/11/2017, anterior ao fato (28/07/2018). Possível erro de digitação na data do termo policial. Verificar nos autos a data correta.", WARNING_BG, WARNING_BORDER)

    # 2.4 Situação prisional
    add_subsection(doc, "2.4 Situação Prisional")

    add_card(doc, "🔴 COSME — Preso desde 29/12/2021", [
        "Unidade: Conjunto Penal de Juazeiro",
        "Tempo preso: ~4 anos e 3 meses (sem condenação em plenário)",
        "EXCESSO DE PRAZO — razoabilidade comprometida",
    ], ERROR_BORDER)

    add_card(doc, "🔴 MARCÉLIO — Preso desde 16/01/2022", [
        "Unidade: Colônia Penal de Simões Filho",
        "Tempo preso: ~4 anos e 2 meses (sem condenação em plenário)",
        "Filho de 14 anos com anemia, sem tratamento por ausência do pai",
        "_'Ele tem só problema de anemia. Não faz [tratamento] porque eu estou preso' — Marcélio, interrogatório_",
    ], ERROR_BORDER)

    add_alert_box(doc, "🔴 CRÍTICO: Ambos os defendidos estão presos há mais de 4 anos SEM julgamento pelo Tribunal do Júri. Configura-se flagrante excesso de prazo. Avaliar impetração de Habeas Corpus com urgência, especialmente considerando a tese de legítima defesa.", ERROR_BG, ERROR_BORDER)

    # 2.5 Histórico DPE
    add_subsection(doc, "2.5 Histórico de Atuação da DPE")
    hdpe = doc.add_table(rows=3, cols=3)
    add_header_row(hdpe, ["Data", "Defensor(a)", "Ação"])
    dpe_data = [
        ("2025", "Dra. Juliane", "Audiência de instrução — sem perguntas às testemunhas"),
        ("2025", "Dr. Rodrigo / Dr. Daniel", "Interrogatório dos defendidos — perguntas estratégicas"),
    ]
    for i, (data, def_, acao) in enumerate(dpe_data):
        add_data_row(hdpe, i+1, [data, def_, acao])
    set_table_borders(hdpe)
    doc.add_paragraph()

    # 2.6 Inventário de documentos
    add_subsection(doc, "2.6 Inventário de Documentos")
    inv = doc.add_table(rows=10, cols=3)
    add_header_row(inv, ["Documento", "Status", "Observação"])
    docs_list = [
        ("Denúncia", "✔", "Nos autos"),
        ("Auto de prisão em flagrante", "❓", "[VERIFICAR]"),
        ("Laudo necroscópico", "❓", "[VERIFICAR] — causa mortis, nº lesões, direção"),
        ("Laudo de local", "✔", "Referenciado — sangue em paredes, teto, chão"),
        ("Depoimento policial Elisangela (2017)", "✔", "Versão 1: Cosme só tomou a faca"),
        ("Depoimento policial Elisangela (2023)", "✔", "Versão 2: ambos esfaquearam — CONTRADIÇÃO"),
        ("Oitivas judiciais (6 testemunhas + interrogatórios)", "✔", "Transcrições em arquivo"),
        ("FAC dos defendidos", "❓", "[VERIFICAR]"),
        ("Antecedentes da suposta vítima (Cláudio)", "⚠️", "Confirmar Marias da Penha / VD"),
    ]
    for i, (doc_name, status, obs) in enumerate(docs_list):
        bg = SUCCESS_BG if status == "✔" else (PENDING_BG if status == "⚠️" else WARNING_BG)
        add_data_row(inv, i+1, [doc_name, status, obs], bg)
    set_table_borders(inv)
    doc.add_paragraph()

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE III — PROVA
    # ========================================================
    add_section_header(doc, "PARTE III — PROVA")

    # 3.1 Quadro de testemunhas
    add_subsection(doc, "3.1 Quadro de Testemunhas")
    qt = doc.add_table(rows=9, cols=5)
    add_header_row(qt, ["Depoente", "Tipo", "Arrolada por", "Presenciou fatos?", "Relevância"])
    testemunhas = [
        ("Elisangela M. Gomes de Souza", "Informante (irmã)", "MP/Defesa", "SIM — presencial", "🔴 CRÍTICA"),
        ("Doraci Mariano Feitosa", "Testemunha", "MP", "NÃO — vizinha abaixo", "⚪ BAIXA"),
        ("Rita Auxiliadora Santana Silva", "Test. EPC (SILC)", "MP", "NÃO — perícia local", "🟡 MÉDIA"),
        ("Driane dos Santos da Silva", "Test. EPC (Escrivã)", "MP", "NÃO — ouviu Elisangela", "🟡 MÉDIA"),
        ("Sandelen Mendes dos Santos", "Testemunha", "MP", "NÃO — ouviu gritos", "⚪ BAIXA"),
        ("Tiago dos Santos", "Testemunha", "MP", "NÃO — chegou depois", "⚪ BAIXA"),
        ("Flávio Cerqueira Santos", "Test. EPC", "MP", "NÃO — não se recorda", "⚪ NULA"),
        ("Cosme (interrogatório)", "Defendido", "—", "SIM — presencial", "🔴 CRÍTICA"),
    ]
    for i, row_data in enumerate(testemunhas):
        add_data_row(qt, i+1, list(row_data), SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(qt)
    doc.add_paragraph()

    # 3.2 Fichas de depoentes
    add_subsection(doc, "3.2 Fichas de Depoentes")

    # --- ELISANGELA ---
    add_card(doc, "ELISANGELA MARIA GOMES DE SOUZA — Informante (irmã dos defendidos, companheira da suposta vítima)", [
        "Fase policial (2017 — Deleg. Geovane): Cosme apenas TOMOU a faca da vítima, não desferiu golpes.",
        "Fase policial (2023 — Deleg. Almir): Mudou versão — ambos esfaquearam Cláudio.",
        "Fase judicial: Narrativa confusa, com dificuldade de lembrar detalhes. Confirma violência de Cláudio.",
        "",
        "🔴 CONTRADIÇÃO CENTRAL:",
    ], ERROR_BORDER)

    add_citation(doc, "Cosme chegou a ser atingido no braço direito. E que a participação de Cosme foi tomar a faca da vítima, ou seja, ele não desferiu nenhum golpe.", "Depoimento policial Elisangela, 01/11/2017 — lido em audiência pela Defesa (Escrivã Driane confirmou)")

    add_citation(doc, "No que eu me recordo ela informou que tinha sido Cosme e Miguel. Miguel começou, ela ligou para Cosme para poder separar a briga e ele deferiu os golpes de faca junto com Marcélio.", "Escrivã Driane sobre o depoimento de Elisangela em 2023")

    add_body(doc, "Em audiência, Elisangela confirma elementos centrais da legítima defesa:")

    add_citation(doc, "Meu marido chegou estressado para a cozinha, pegou a faca e disse que ia [...] meu marido entrou dentro de casa, pegou a faca e entrou para dentro do quarto.", "Elisangela, audiência judicial")

    add_citation(doc, "Ele me batia todos os dias sem eu fazer nada, me traia, botava ele em casa [...] Rapaz eu passei mal com o senhor, com o homem", "Elisangela, audiência judicial — sobre violência de Cláudio")

    add_citation(doc, "Eu tinha medo, não vou mentir, foi ameaçada, eu tinha medo.", "Elisangela, audiência judicial — sobre não registrar ocorrências")

    add_citation(doc, "Ele falou, se eu não abrisse a porta, ele ia me matar.", "Elisangela, audiência judicial — Cláudio ameaçou matá-la no dia dos fatos")

    add_citation(doc, "Eu chamei [Cosme] para ver, você tirava a maçã dedicada para evitar confusão.", "Elisangela, audiência judicial — chamou Cosme para SEPARAR")

    add_citation(doc, "O Cosme bolou com ele para pôr uma faca.", "Elisangela, audiência judicial — Cosme interveio para tomar a faca")

    add_citation(doc, "Só tinha uma [faca]. E até sumiu, essa faca sumiu de lá. Não sei quem pegou.", "Elisangela, audiência judicial — apenas UMA faca no local")

    add_body(doc, "Avaliação de credibilidade: COMPROMETIDA pela contradição entre versões policiais. Porém, em ambas as versões, Cláudio é o agressor inicial. A mudança prejudica especificamente a posição de COSME (de não-agressor para co-agressor).", bold=True)

    doc.add_paragraph()

    # --- DORACI ---
    add_card(doc, "DORACI MARIANO FEITOSA — Vizinha do andar de baixo", [
        "Não presenciou os fatos. Fechou janela e aumentou volume da TV.",
        "No dia seguinte, perguntou a Elisangela o que aconteceu.",
        "Não tem relação com os acusados.",
    ], MUTED)

    add_citation(doc, "Para ser sincera com o senhor [...] eu não vim para aqui para mentir [...] quando o dia da briga que começou quebra-pau lá para cima, eu tenho um problema de saúde, aí eu fechei minha janela, aumentei o volume da minha televisão e fiquei quietinha.", "Doraci, audiência judicial")

    add_citation(doc, "Ela me disse assim: olha Nadora, eles estavam na roça, quando chegaram aqui já vieram brigando de lá, da roça para casa. Quando chegou aqui, aí eles começaram a discutir, por essa discussão terminou matando ele.", "Doraci sobre o que Elisangela contou — hearsay")

    add_citation(doc, "Ela me disse que quem matou foi o Marcelo.", "Doraci sobre relato de Elisangela")

    add_body(doc, "Avaliação: Testemunha de ouvir dizer (hearsay). Não presenciou nada. Seu relato apenas repete o que Elisangela disse, que já mudou de versão.", italic=True)
    doc.add_paragraph()

    # --- RITA (EPC) ---
    add_card(doc, "RITA AUXILIADORA SANTANA SILVA — EPC/SILC (perícia de local)", [
        "Lotada na 13ª DP (Cajazeiras). Participou apenas do SILC (perícia de local).",
        "Não se recorda bem da ocorrência — muitas diligências.",
        "Relata que encontrou 'muito, muito sangue' no apartamento.",
    ], MUTED)

    add_citation(doc, "Quando nós chegamos, nós vimos que tinha sangue por todo canto. Paredes, teto, chão. Muito, muito sangue.", "Rita, audiência judicial — sobre o local")

    add_citation(doc, "Ela [Elisangela] relatou sim [...] Parece que os três trabalhavam em alguma coisa de roça [...] tinham tido algum desentendimento entre eles.", "Rita sobre relato de Elisangela no local")

    add_citation(doc, "Ela falou que eles fugiram por alguma janela, um basculante, que ela chegou a mostrar em um dos quartos.", "Rita — sobre a fuga dos defendidos")

    add_citation(doc, "Não houve nenhum relato sobre o motivo.", "Rita — ninguém soube explicar o motivo da violência")

    add_body(doc, "Avaliação: Testemunha técnica. Não acrescenta autoria — apenas descreve o local. Confirma que Elisangela não soube explicar o motivo.", italic=True)
    doc.add_paragraph()

    # --- DRIANE (EPC) ---
    add_card(doc, "DRIANE DOS SANTOS DA SILVA — EPC/Escrivã (Deleg. Almir)", [
        "Participou dos depoimentos de Elisangela com o delegado Almir (2023).",
        "NÃO participou do primeiro depoimento (2017 — Escrivã Tainara).",
        "CONFIRMOU a contradição entre as versões de Elisangela.",
    ], WARNING_BORDER)

    add_citation(doc, "Pelos depoimentos, a companheira da vítima informa que o seu irmão Miguel e o seu companheiro Cláudio tiveram uma briga. Eles trabalhavam em uma roça, num sítio chamado Seu Brandão. [...] Ela ligou para o outro irmão dela, chamado Cosme. Para separar a briga, porém quando chegou lá, ao invés de ele separar, ele também golpeou o companheiro dela com faca.", "Driane, audiência judicial")

    add_citation(doc, "No que eu me recordo ela informou que tinha sido Cosme e Miguel. Miguel começou, ela ligou para Cosme para poder separar a briga e ele deferiu os golpes de faca junto com Marcélio.", "Driane — sobre versão de 2023 de Elisangela")

    add_body(doc, "Avaliação: Testemunha crucial para demonstrar a CONTRADIÇÃO de Elisangela. Confirmou que leu o depoimento de 2017 (Tainara) onde Cosme NÃO desferiu golpes. Não houve confronto com Elisangela sobre essa contradição.", italic=True, bold=True)
    doc.add_paragraph()

    # --- SANDELEN ---
    add_card(doc, "SANDELEN MENDES DOS SANTOS — Vizinha", [
        "Recém-chegada ao condomínio. Não conhecia os envolvidos.",
        "Ouviu gritos de Elisangela pedindo socorro. Não entrou na casa.",
    ], MUTED)

    add_citation(doc, "No dia do ocorrido, eu estava em casa, na parte da cozinha, quando eu só ouvi os gritos da irmã do acusado e a única coisa que eu presenciei foi as gritarias.", "Sandelen, audiência judicial")

    add_citation(doc, "Quando ela saiu da casa, ela pedindo por socorro, já tinha gente já ligando para a polícia.", "Sandelen — sobre Elisangela pedindo socorro")

    add_body(doc, "Avaliação: Nenhum conhecimento direto. Apenas reforça que Elisangela pediu socorro.", italic=True)
    doc.add_paragraph()

    # --- TIAGO ---
    add_card(doc, "TIAGO DOS SANTOS — Vizinho", [
        "Não estava em casa no momento dos fatos. Chegou depois.",
        "Viu movimentação e soube que 'Negão' estava no chão sangrando.",
    ], MUTED)

    add_citation(doc, "Na noite que aconteceu isso [...] eu não estava na residência. [...] Aí o pessoal que estava na rua falou, não, houve uma briga aqui [...] Negão foi, está ali no chão ali sangrando.", "Tiago, audiência judicial")

    add_body(doc, "Avaliação: Irrelevante probatoriamente. Não presenciou, não ouviu relatos detalhados.", italic=True)
    doc.add_paragraph()

    # --- FLAVIO ---
    add_card(doc, "FLÁVIO CERQUEIRA SANTOS — EPC/Plantão", [
        "Não se recorda absolutamente nada da ocorrência.",
        "Trabalha em plantões com múltiplos homicídios por dia.",
    ], MUTED)

    add_citation(doc, "Não, Excelência. Não se recorda nada dos fatos [...] às vezes são tantas situações, né, com uma certa feita me fiz oito homicídio num dia só, e aí eu não tenho como lembrar.", "Flávio, audiência judicial")

    add_body(doc, "Avaliação: Contribuição probatória NULA.", italic=True)
    doc.add_paragraph()

    # --- INTERROGATÓRIO COSME ---
    add_card(doc, "INTERROGATÓRIO — Cosme Miguel Gomes de Souza (Defendido 1)", [
        "Respondeu apenas às perguntas da defesa (orientação técnica).",
        "Versão coerente e detalhada.",
    ], ACCENT)

    add_citation(doc, "Na verdade, doutor, ele já vinha ameaçando a gente já havia mais de cinco anos, a verdade é essa.", "Cosme, interrogatório")

    add_citation(doc, "Eu estava trabalhando na roça. Chegou os dois [...] lá eles já tinham brigado os dois [...] meu irmão mais meu ex-cunhado.", "Cosme — Marcélio e Cláudio brigaram antes na roça")

    add_citation(doc, "Eu só fui mesmo para o meu irmão tirar em roupa, quando cheguei lá, o que eu estava no banheiro aconteceu, ele já trancou meu irmão dentro do quarto, já foi para desferir a faca do meu irmão, minha irmã começou a gritar, me chamou às pressas.", "Cosme — dinâmica dos fatos")

    add_citation(doc, "Quando eu invadi o quarto, ele já estava com a faca, levei o golpe primeiro no braço, e aconteceu, meu irmão tomou a faca e aconteceu o que tinha de acontecer.", "Cosme — foi atingido primeiro, Marcélio tomou a faca")

    add_citation(doc, "Foi legítimo a defesa, meu irmão.", "Cosme — afirma legítima defesa")

    add_citation(doc, "Ele ameaçava de matar minha irmã, como ameaçava de matar a gente, como já tinha tentado várias vezes contra mim e não conseguiu. Eu me afastei, fui pra longe pra não causar problema com ele.", "Cosme — histórico de ameaças")

    add_citation(doc, "Ele também já estava com outra mulher dentro da casa de minha irmã, morando com outra mulher, com minha irmã, obrigando minha irmã a aceitar.", "Cosme — contexto de violência doméstica")

    add_body(doc, "Avaliação: Versão clara, coerente com os demais elementos (briga prévia na roça, Cláudio como agressor inicial, Cosme atingido no braço). Sustenta legítima defesa.", italic=True, bold=True)
    doc.add_paragraph()

    # --- INTERROGATÓRIO MARCÉLIO ---
    add_card(doc, "INTERROGATÓRIO — Marcélio Miguel de Souza (Defendido 2)", [
        "Respondeu apenas às perguntas da defesa.",
        "Versão coerente com a de Cosme.",
    ], ACCENT)

    add_citation(doc, "Isso começou quando nós estávamos bebendo na roça porque nós trabalhávamos na mesma roça e morávamos na mesma casa. Tanto que ele era meu cunhado. Aí por causa de uma simples ligação no meu celular de uma mulher, aí ele foi atender, aí começou todo o rebuliço.", "Marcélio — o estopim na roça")

    add_citation(doc, "Quando eu cheguei em casa, ele não estava deixando eu pegar minhas roupas, e já estava com uma faca na mão.", "Marcélio — Cláudio armado com faca esperando")

    add_citation(doc, "Minha irmã dentro do quarto, se ela saísse ele ia fazer algo com ela, como ele estava jurando de matar ela.", "Marcélio — Cláudio ameaçava Elisangela")

    add_citation(doc, "Minha irmã gritou para o meu irmão Cosme, para poder separar, para poder tirar. Aí foi onde ele desferiu um golpe de faca no meu irmão, que a faca caiu no chão e aí foi onde aconteceu o fato.", "Marcélio — Cláudio golpeou Cosme, faca caiu, Marcélio pegou")

    add_citation(doc, "Foi eu. Quando eu vi aquele sangue foi.", "Marcélio — admite que desferiu facadas depois que a faca caiu")

    add_citation(doc, "Ele batia na minha irmã, realmente, que ele já tinha um histórico [...] Inclusive, ele botou uma mulher grávida pra morar junto com minha irmã.", "Marcélio — violência de Cláudio contra Elisangela")

    add_citation(doc, "Nós já tínhamos discutido por causa disso já. Mas sempre a discussão nunca elevou ao fato que levou dessa forma.", "Marcélio — conflitos anteriores com Cláudio")

    add_body(doc, "Avaliação: Versão coerente com a de Cosme. Admite ter desferido os golpes, mas após ter sido atacado primeiro com faca. Contexto de violência prévia e ameaças sustenta a legítima defesa.", italic=True, bold=True)
    doc.add_paragraph()

    # 3.3 Tabela comparativa
    add_subsection(doc, "3.3 Tabela Comparativa de Depoimentos")
    comp = doc.add_table(rows=8, cols=5)
    add_header_row(comp, ["Ponto", "Elisangela (2017)", "Elisangela (2023/Juízo)", "Cosme", "Marcélio"])
    comparativos = [
        ("Quem pegou a faca primeiro?", "Cláudio", "Cláudio", "Cláudio", "Cláudio"),
        ("Cláudio atacou primeiro?", "Sim (implícito)", "Sim", "Sim — golpeou Cosme no braço", "Sim — estava armado esperando"),
        ("Cosme desferiu golpes?", "NÃO — só tomou a faca", "SIM — esfaqueou também", "Não menciona", "Não — Cosme foi atingido"),
        ("Marcélio desferiu golpes?", "Sim", "Sim", "Marcélio tomou a faca", "Sim — admite"),
        ("Quantas facas?", "—", "1 faca só", "1 faca (de Cláudio)", "1 faca (caiu no chão)"),
        ("Elisangela chamou Cosme?", "Sim — para separar", "Sim — para separar", "Sim — chamou às pressas", "Sim — irmã gritou"),
        ("Fuga pela janela?", "Sim", "Sim", "—", "Sim — pularam pela janela"),
    ]
    for i, row_data in enumerate(comparativos):
        add_data_row(comp, i+1, list(row_data), SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(comp)
    doc.add_paragraph()

    add_alert_box(doc, "🔴 CONTRADIÇÃO-CHAVE: Na versão de 2017, Cosme NÃO desferiu golpes — apenas tomou a faca. Na versão de 2023, Cosme também esfaqueou. Esta mudança nunca foi confrontada formalmente com Elisangela. A escrivã Driane confirmou a existência das duas versões contraditórias.", ERROR_BG, ERROR_BORDER)

    # 3.4 Provas materiais e lacunas
    add_subsection(doc, "3.4 Provas Materiais e Lacunas")

    pmla = doc.add_table(rows=7, cols=3)
    add_header_row(pmla, ["Prova", "Existe?", "Observação"])
    provas = [
        ("Laudo necroscópico", "❓", "[VERIFICAR] Causa mortis, nº de lesões, instrumento"),
        ("Perícia de local", "✔", "Muito sangue — paredes, teto, chão. Sangue na cortina/janela (fuga)"),
        ("Arma do crime (faca/peixeira)", "❌", "Não apreendida — 'sumiu' segundo Elisangela"),
        ("Exame de corpo de delito (Cosme)", "❌", "Cosme foi ferido no braço — não há laudo"),
        ("Exame de corpo de delito (Marcélio)", "❌", "Marcélio foi cortado — não há laudo"),
        ("Antecedentes Cláudio (VD/ameaças)", "⚠️", "Mencionados por testemunhas — necessário juntar FAC/BOs"),
    ]
    for i, (prova, existe, obs) in enumerate(provas):
        bg = SUCCESS_BG if existe == "✔" else (ERROR_BG if existe == "❌" else PENDING_BG)
        add_data_row(pmla, i+1, [prova, existe, obs], bg)
    set_table_borders(pmla)
    doc.add_paragraph()

    # 3.5 Fragilidades probatórias
    add_subsection(doc, "3.5 Fragilidades Probatórias")

    fragilidades = [
        ("🔴 CRÍTICO", "1. Contradição de Elisangela entre versões policiais (2017 vs 2023)", "Na primeira versão, Cosme NÃO esfaqueou ninguém. Na segunda, participou ativamente. A mudança nunca foi confrontada."),
        ("🔴 CRÍTICO", "2. Nenhuma testemunha presencial além de Elisangela (comprometida)", "Doraci, Sandelen, Tiago: nada viram. Rita e Driane: só ouviram relato de Elisangela. Flávio: não lembra de nada."),
        ("🟡 ALTO", "3. Arma do crime não apreendida", "A faca 'sumiu' — não há perícia no instrumento, não se sabe se havia mais de uma."),
        ("🟡 ALTO", "4. Ausência de exame de corpo de delito nos defendidos", "Cosme e Marcélio foram feridos por Cláudio, mas não há laudo comprovando suas lesões defensivas."),
        ("⚪ MÉDIO", "5. Histórico de violência de Cláudio não documentado formalmente nos autos", "Todas as testemunhas mencionam, mas falta juntar BOs, FAC e processos anteriores de Cláudio."),
        ("⚪ MÉDIO", "6. Elisangela em estado de choque no dia — confiabilidade do primeiro relato", "Estava grávida, em estado de choque, sozinha na delegacia sem advogado."),
    ]
    for severity, title, detail in fragilidades:
        add_card(doc, f"{severity} — {title}", [detail], ERROR_BORDER if "CRÍTICO" in severity else (WARNING_BORDER if "ALTO" in severity else MUTED))

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE IV — ESTRATÉGIA
    # ========================================================
    add_section_header(doc, "PARTE IV — ESTRATÉGIA")

    # 4.1 Teses defensivas
    add_subsection(doc, "4.1 Teses Defensivas")

    # Tese 1
    add_card(doc, "TESE 1 — LEGÍTIMA DEFESA PRÓPRIA E DE TERCEIRO", [
        "Viabilidade: ■■■■□ ALTA",
        "",
        "Fundamento: Art. 25 do CP — reação a agressão injusta, atual e iminente.",
        "",
        "Para Marcélio: Legítima defesa PRÓPRIA. Cláudio pegou a faca, trancou-se esperando,",
        "e atacou Marcélio quando este entrou para pegar roupas. Marcélio foi cortado no braço.",
        "Ao tomar a faca, reagiu contra o agressor.",
        "",
        "Para Cosme: Legítima defesa DE TERCEIRO. Chamado por Elisangela para separar a briga.",
        "Ao entrar no quarto, foi golpeado no braço por Cláudio. Interveio para proteger",
        "o irmão que estava sendo atacado com faca.",
        "",
        "Elementos favoráveis:",
        "• Cláudio foi quem pegou a faca primeiro (todas as versões concordam)",
        "• Cláudio ameaçava a família há 5+ anos",
        "• Cláudio estava armado esperando atrás da porta",
        "• Cláudio ameaçou matar Elisangela ('se eu não abrisse a porta, ele ia me matar')",
        "• Cosme e Marcélio foram feridos (lesões defensivas)",
        "• Havia apenas UMA faca — a de Cláudio",
        "• Na primeira versão policial, Cosme sequer desferiu golpes",
        "",
        "Riscos: MP argumentará excesso (múltiplas facadas). Necessário demonstrar",
        "que o contexto de terror (5 anos de ameaças) e a dinâmica da briga",
        "justificam a intensidade da reação.",
    ], ACCENT)

    # Tese 2
    add_card(doc, "TESE 2 (Subsidiária) — EXCESSO CULPOSO NA LEGÍTIMA DEFESA", [
        "Viabilidade: ■■■□□ MÉDIA",
        "",
        "Fundamento: Art. 23, parágrafo único, do CP.",
        "Se os jurados reconhecerem legítima defesa mas entenderem que houve excesso,",
        "a desclassificação para homicídio culposo reduz significativamente a pena.",
        "",
        "Argumento: No calor da luta corporal, com sangue e adrenalina, os defendidos",
        "ultrapassaram os limites da moderação sem intenção de matar, apenas tentando",
        "neutralizar a ameaça.",
        "",
        "Risco: Se o júri negar a legítima defesa, esta tese cai junto.",
    ], WARNING_BORDER)

    # Tese 3
    add_card(doc, "TESE 3 (Subsidiária) — HOMICÍDIO PRIVILEGIADO POR VIOLENTA EMOÇÃO", [
        "Viabilidade: ■■■□□ MÉDIA",
        "",
        "Fundamento: Art. 121, §1º, do CP — domínio de violenta emoção, logo em seguida",
        "a injusta provocação da vítima.",
        "",
        "Argumento: Após 5+ anos de violência, ameaças e humilhações contra a família,",
        "ver o cunhado armado com faca, ameaçando matar a irmã e o irmão, gerou",
        "uma reação de violenta emoção compreensível.",
        "",
        "Efeito: Redução de 1/6 a 1/3 da pena.",
    ], WARNING_BORDER)

    # Tese 4
    add_card(doc, "TESE 4 (Específica para Cosme) — PARTICIPAÇÃO DE MENOR IMPORTÂNCIA", [
        "Viabilidade: ■■■□□ MÉDIA",
        "",
        "Fundamento: Art. 29, §1º, do CP.",
        "",
        "Argumento: Na primeira versão de Elisangela (2017), Cosme NÃO desferiu golpes —",
        "apenas tomou a faca de Cláudio. Se os jurados acreditarem nesta versão (mais",
        "próxima dos fatos, sem contaminação), a participação de Cosme foi mínima.",
        "",
        "Efeito: Redução de 1/6 a 1/3 da pena.",
        "",
        "Risco: Versão de 2023 contradiz (ambos esfaquearam).",
    ], WARNING_BORDER)

    # 4.2 Narrativa defensiva
    add_subsection(doc, "4.2 Narrativa Defensiva")
    add_body(doc, "Cláudio de Jesus Silva era um homem violento que por mais de cinco anos aterrorizou sua companheira Elisangela e sua família com agressões físicas diárias, ameaças de morte e humilhações — chegando a instalar outra mulher grávida dentro da casa de Elisangela, obrigando-a a conviver com a situação sob ameaça. No dia 28 de julho de 2018, após uma discussão com Marcélio na roça onde ambos trabalhavam, Cláudio voltou para casa com o rosto ferido, pegou uma faca grande (peixeira) na cozinha, e se posicionou armado atrás da porta do quarto, esperando Marcélio, que apenas queria recolher suas roupas e ir embora. Quando Marcélio entrou, Cláudio partiu para cima dele com a faca, cortando seu braço. Elisangela, desesperada — grávida e com três crianças pequenas dentro da casa — gritou por Cosme, que estava no banheiro, pedindo que separasse a briga. Ao tentar intervir, Cosme também foi atingido pela faca no braço. Na luta que se seguiu, a faca caiu e Marcélio a pegou, reagindo contra o agressor para salvar a si mesmo, ao irmão e à irmã que Cláudio jurava matar. Não houve crime — houve sobrevivência.", bold=True)

    # 4.3 Projeção de dosimetria
    add_subsection(doc, "4.3 Projeção de Dosimetria (em caso de condenação)")
    dosi = doc.add_table(rows=9, cols=3)
    add_header_row(dosi, ["Etapa", "Marcélio", "Cosme"])
    dosi_data = [
        ("Pena base (art. 121 caput)", "6 anos", "6 anos"),
        ("Circunstâncias judiciais (art. 59)", "Mínimo — primário, bons antecedentes", "Mínimo — [VERIFICAR antecedentes]"),
        ("Atenuantes", "Menoridade relativa (< 21 à época)", "—"),
        ("Agravantes", "—", "—"),
        ("Causas de diminuição", "Violenta emoção / Excesso culposo / Participação menor", "Participação menor (versão 2017)"),
        ("Pena provável (condenação)", "4-5 anos", "4-5 anos"),
        ("Detração", "~4 anos e 2 meses (desde 16/01/2022)", "~4 anos e 3 meses (desde 29/12/2021)"),
        ("Regime inicial", "Aberto (se < 4 anos) ou Semiaberto", "Aberto ou Semiaberto"),
    ]
    for i, row_data in enumerate(dosi_data):
        add_data_row(dosi, i+1, list(row_data), SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(dosi)
    doc.add_paragraph()

    add_alert_box(doc, "⚠ PONTO ESTRATÉGICO: Mesmo em caso de condenação, a detração (4+ anos) possivelmente cobriria ou quase cobriria a pena, especialmente com as causas de diminuição. Os defendidos poderiam ser colocados em liberdade imediata.", WARNING_BG, WARNING_BORDER)

    # 4.4 Matriz de riscos
    add_subsection(doc, "4.4 Matriz de Riscos")
    risk = doc.add_table(rows=4, cols=4)
    add_header_row(risk, ["Risco", "Probabilidade", "Impacto", "Mitigação"])
    risks = [
        ("Júri não aceitar legítima defesa", "MÉDIA", "ALTO", "Construir narrativa de terror familiar. Explorar contradição"),
        ("MP enfatizar múltiplas facadas", "ALTA", "MÉDIO", "Contexto de luta corporal, adrenalina, 1 faca só"),
        ("Elisangela depor contra irmãos", "BAIXA", "ALTO", "Usar contradição a favor — versão de 2017"),
    ]
    for i, row_data in enumerate(risks):
        add_data_row(risk, i+1, list(row_data), SUBTLE_BG if i % 2 == 0 else WHITE)
    set_table_borders(risk)
    doc.add_paragraph()

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE V — PLENÁRIO
    # ========================================================
    add_section_header(doc, "PARTE V — PLENÁRIO")

    # 5.1 Roteiro de atendimento prévio
    add_subsection(doc, "5.1 Roteiro de Atendimento Prévio")
    checklist_atend = [
        "☐ Visitar Cosme no Conjunto Penal de Juazeiro antes do plenário",
        "☐ Visitar Marcélio na Colônia Penal de Simões Filho antes do plenário",
        "☐ Repassar a versão com cada um — garantir coerência entre as narrativas",
        "☐ Orientar sobre postura em plenário (olhar nos olhos dos jurados, falar pausado)",
        "☐ Verificar se têm roupas adequadas para o plenário",
        "☐ Recolher informações sobre comportamento carcerário (atestado)",
        "☐ Verificar se Marcélio quer falar sobre o filho com anemia",
        "☐ Confirmar se querem declarar em plenário (recomendável — versão é favorável)",
        "☐ Solicitar providências familiares (fotos, documentos do filho de Marcélio)",
        "☐ Verificar contato com Elisangela — posição atual dela",
    ]
    for item in checklist_atend:
        add_body(doc, item)

    # 5.2 Protocolo do dia
    add_subsection(doc, "5.2 Protocolo do Dia")
    protocolo = [
        "1. Requisição dos defendidos com antecedência mínima de 72h",
        "2. Verificar presença das testemunhas (Elisangela é crucial)",
        "3. Sorteio e recusa de jurados — atenção ao perfil",
        "4. Oitiva de testemunhas de acusação",
        "5. Oitiva de testemunhas de defesa (se houver)",
        "6. Interrogatório dos defendidos (ambos — orientar para declarar)",
        "7. Debates: acusação (1h30) → defesa (1h30) → réplica (1h) → tréplica (1h)",
        "8. Quesitação",
        "9. Votação",
    ]
    for item in protocolo:
        add_body(doc, item)

    # 5.3 Perguntas estratégicas por depoente
    add_subsection(doc, "5.3 Perguntas Estratégicas por Depoente")

    add_card(doc, "ELISANGELA — Perguntas da Defesa", [
        "Objetivo: Reforçar que Cláudio era o agressor / explorar contradição",
        "",
        "1. A senhora pode contar aos jurados como era a convivência com Cláudio? Ele era violento?",
        "   → Objetivo: Estabelecer o padrão de violência",
        "",
        "2. A senhora tinha medo de Cláudio? Ele já ameaçou matar a senhora ou seus irmãos?",
        "   → Objetivo: Demonstrar o contexto de terror",
        "",
        "3. No dia dos fatos, quem pegou a faca primeiro?",
        "   → Objetivo: Confirmar que Cláudio foi o agressor inicial (ponto unânime)",
        "",
        "4. A senhora chamou Cosme para quê? Para separar a briga ou para agredir Cláudio?",
        "   → Objetivo: Demonstrar que a intenção era pacificar",
        "",
        "5. Na delegacia, em 2017, a senhora disse que Cosme apenas tomou a faca sem desferir golpes. Isso é verdade?",
        "   → Objetivo: Resgatar a primeira versão favorável a Cosme",
        "",
        "6. Quantas facas havia na cena? Apenas a faca que Cláudio pegou?",
        "   → Objetivo: Confirmar que era 1 faca — a de Cláudio",
    ], ACCENT)

    add_card(doc, "DORACI — Perguntas da Defesa", [
        "Objetivo: Minimizar o impacto do hearsay / reforçar desconhecimento",
        "",
        "1. A senhora não viu os fatos, correto? Apenas soube depois?",
        "2. A senhora conhecia o temperamento de Cláudio?",
    ], MUTED)

    add_card(doc, "DRIANE (Escrivã) — Perguntas da Defesa", [
        "Objetivo: EXPLORAR a contradição entre as versões de Elisangela",
        "",
        "1. A senhora leu o primeiro depoimento de Elisangela, de 2017, correto?",
        "   → Objetivo: Confirmar acesso à versão 1",
        "",
        "2. Nesse depoimento, consta que Cosme apenas TOMOU a faca, sem desferir golpes. A senhora se recorda disso?",
        "   → Objetivo: Ratificar a contradição em plenário",
        "",
        "3. Quando Elisangela foi ouvida em 2023, ela foi confrontada com essa contradição?",
        "   → Objetivo: Demonstrar que a mudança não foi questionada pela polícia",
        "",
        "4. Em seu entender, a primeira versão — mais próxima dos fatos — é mais confiável?",
        "   → Objetivo: Semear dúvida nos jurados (cuidado: objeção do MP)",
    ], WARNING_BORDER)

    # 5.4 Orientação aos defendidos
    add_subsection(doc, "5.4 Orientação aos Defendidos")

    add_card(doc, "COSME — Orientação para Plenário", [
        "✔ DECLARAR em plenário (versão é favorável)",
        "✔ Enfatizar: foi chamado para SEPARAR, foi atingido no braço, tentou ajudar",
        "✔ Mencionar os 5+ anos de ameaças e violência contra a irmã",
        "✔ Demonstrar arrependimento pelo resultado, mas firmar que agiu para proteger",
        "✔ Tom: calmo, firme, respeitoso",
        "✘ NÃO brigar com o MP, NÃO alterar a voz",
        "✘ NÃO dizer 'ele mereceu' ou qualquer coisa que sugira intenção de matar",
    ], ACCENT)

    add_card(doc, "MARCÉLIO — Orientação para Plenário", [
        "✔ DECLARAR em plenário (versão é favorável)",
        "✔ Enfatizar: foi atacado com faca, cortado no braço, reagiu em desespero",
        "✔ Mencionar que só queria pegar roupas e ir embora",
        "✔ Falar sobre o filho de 14 anos com anemia, sem tratamento",
        "✔ Tom: calmo, honesto, vulnerável (é jovem, tem filho doente)",
        "✘ NÃO minimizar — admitir que desferiu os golpes, mas em defesa",
        "✘ NÃO culpar Elisangela pela mudança de versão",
    ], ACCENT)

    # 5.5 Requerimentos orais prontos
    add_subsection(doc, "5.5 Requerimentos Orais Prontos")

    add_card(doc, "REQUERIMENTO 1 — Leitura de peça", [
        "Requeiro a Vossa Excelência que seja lido em plenário o depoimento prestado por",
        "Elisangela Maria Gomes de Souza perante o Delegado Geovane, em novembro de 2017,",
        "no qual consta expressamente que 'a participação de Cosme foi tomar a faca da",
        "vítima, ou seja, ele não desferiu nenhum golpe', nos termos do art. 473, §3º, do CPP.",
    ], ACCENT)

    add_card(doc, "REQUERIMENTO 2 — Diligência", [
        "Requeiro seja juntada aos autos a Folha de Antecedentes Criminais de Cláudio de",
        "Jesus Silva, a fim de demonstrar aos jurados o histórico de violência da suposta",
        "vítima, elemento essencial para a compreensão do contexto de legítima defesa.",
    ], ACCENT)

    add_card(doc, "REQUERIMENTO 3 — Excesso de prazo", [
        "Requeiro que conste em ata a manifestação desta Defesa quanto ao flagrante excesso",
        "de prazo da prisão preventiva dos defendidos, presos há mais de 4 anos sem",
        "julgamento pelo Tribunal do Júri, em violação ao art. 5º, LXXVIII, da CF/88.",
    ], ERROR_BORDER)

    # 5.6 Quesitação
    add_subsection(doc, "5.6 Quesitação Prevista e Estratégia")

    add_card(doc, "QUESITO 1 — Materialidade", [
        "'O acusado [nome] matou a vítima Cláudio de Jesus Silva conforme descrito na pronúncia?'",
        "",
        "Estratégia: NÃO contestar materialidade. O fato ocorreu. Concentrar na causa.",
        "Voto esperado: SIM",
    ], MUTED)

    add_card(doc, "QUESITO 2 — Autoria/Participação", [
        "'O acusado [nome] concorreu para a morte da vítima?'",
        "",
        "Estratégia para COSME: Explorar versão de 2017 (não desferiu golpes).",
        "Se possível, buscar NÃO neste quesito para Cosme.",
        "Estratégia para MARCÉLIO: SIM — mas defender nos quesitos seguintes.",
    ], WARNING_BORDER)

    add_card(doc, "QUESITO 3 — Absolvição genérica (art. 483, §2º, CPP)", [
        "'O jurado absolve o acusado?'",
        "",
        "🔴 QUESITO DECISIVO — Este é o campo de batalha.",
        "Estratégia: Pedir SIM com base em LEGÍTIMA DEFESA.",
        "Argumentar: clemência, justiça, 4 anos preso, contexto de terror familiar.",
        "Os jurados podem absolver por QUALQUER razão — não precisam fundamentar.",
    ], ACCENT)

    add_card(doc, "QUESITOS SUBSIDIÁRIOS (se não absolvidos)", [
        "4. 'O acusado agiu em legítima defesa?' → SIM (defesa técnica)",
        "5. 'Houve excesso na legítima defesa?' → Subsidiário",
        "6. 'O acusado cometeu o crime sob domínio de violenta emoção?' → Subsidiário",
        "7. 'A participação do acusado [Cosme] foi de menor importância?' → Para Cosme",
    ], WARNING_BORDER)

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE VI — CENÁRIOS
    # ========================================================
    add_section_header(doc, "PARTE VI — CENÁRIOS")

    # 6.1 Absolvição
    add_subsection(doc, "6.1 Cenário: Absolvição (ambos)")
    add_card(doc, "Absolvição pelo Conselho de Sentença", [
        "Probabilidade: ■■■□□ MÉDIA-ALTA (se bem sustentada)",
        "",
        "Providências imediatas:",
        "☐ Requerer expedição de alvará de soltura em plenário",
        "☐ Comunicar às unidades prisionais",
        "☐ Orientar sobre recursos do MP (apelação — prazo de 5 dias)",
        "☐ Verificar necessidade de medidas cautelares diversas",
    ], SUCCESS_BG)

    # 6.2 Condenação
    add_subsection(doc, "6.2 Cenário: Condenação")
    add_card(doc, "Condenação pelo Conselho de Sentença", [
        "Pena esperada (com atenuantes): 4-6 anos",
        "Detração: 4+ anos já cumpridos",
        "",
        "Se pena ≤ tempo preso: Liberdade imediata",
        "Se pena > tempo preso: Regime aberto ou semiaberto (pouco tempo restante)",
        "",
        "Providências:",
        "☐ Requerer detração em sentença",
        "☐ Requerer regime inicial compatível",
        "☐ Interpor apelação (art. 593, III, 'd', CPP) — decisão contrária à prova",
        "☐ Requerer liberdade provisória até julgamento da apelação",
    ], WARNING_BG)

    # 6.3 Desclassificação
    add_subsection(doc, "6.3 Cenário: Desclassificação")
    add_card(doc, "Desclassificação para homicídio culposo", [
        "Se reconhecido excesso culposo na legítima defesa:",
        "Pena: 1-3 anos de detenção",
        "Com detração de 4+ anos: EXCEDIDO — liberdade imediata",
        "",
        "Se desclassificação para lesão corporal seguida de morte (art. 129, §3º):",
        "Pena: 4-12 anos",
        "Competência: Juiz singular (não mais o júri)",
    ], SUBTLE_BG)

    # 6.4 Contingências
    add_subsection(doc, "6.4 Contingências")
    contingencias = [
        ("Elisangela não comparece", "Requerer sua dispensa ou, se relevante, condução coercitiva. Usar depoimentos policiais (leitura em plenário)."),
        ("Jurado com ligação com MP/Polícia", "Exercer recusa peremptória. Verificar perfil durante sorteio."),
        ("MP pede desaforamento", "Contestar — não há comprometimento da imparcialidade em Camaçari."),
        ("Defendidos se contradizem em plenário", "Orientar previamente. Se ocorrer, focar na versão mais favorável e na prova documental."),
    ]
    for titulo, detalhe in contingencias:
        add_card(doc, f"Contingência: {titulo}", [detalhe], WARNING_BORDER)

    doc.add_page_break()
    add_separator(doc)

    # ========================================================
    # PARTE VII — PROVIDÊNCIAS
    # ========================================================
    add_section_header(doc, "PARTE VII — PROVIDÊNCIAS")

    # 7.1 Urgentes
    add_subsection(doc, "7.1 Providências Urgentes")
    urgentes = [
        "☐ Avaliar impetração de HABEAS CORPUS por excesso de prazo (4+ anos preso sem plenário)",
        "☐ Juntar FAC de Cláudio de Jesus Silva (antecedentes por VD e ameaças)",
        "☐ Juntar BOs anteriores de Elisangela contra Cláudio [VERIFICAR existência]",
        "☐ Solicitar atestado de comportamento carcerário de AMBOS os defendidos",
        "☐ Visitar ambos os defendidos nas unidades prisionais para preparação",
        "☐ Verificar laudo necroscópico — causa mortis, nº e localização das lesões",
        "☐ Verificar se há exame de corpo de delito nos defendidos (lesões defensivas)",
        "☐ Confirmar data correta do depoimento de 2017 de Elisangela (possível erro)",
    ]
    for item in urgentes:
        add_body(doc, item)

    # 7.2 Em plenário
    add_subsection(doc, "7.2 Providências em Plenário")
    plenario_items = [
        "☐ Requerer leitura do depoimento de Elisangela de 2017 (Cosme não esfaqueou)",
        "☐ Explorar contradição de Elisangela na arguição",
        "☐ Requerer juntada da FAC de Cláudio",
        "☐ Registrar excesso de prazo em ata",
        "☐ Utilizar narrativa do terror familiar (5+ anos de violência)",
        "☐ Enfatizar que Cláudio era o agressor inicial (ponto unânime em todas as versões)",
        "☐ Pedir absolvição genérica com base em legítima defesa e clemência",
        "☐ Subsidiariamente, sustentar excesso culposo e homicídio privilegiado",
    ]
    for item in plenario_items:
        add_body(doc, item)

    # 7.3 Pós-plenário
    add_subsection(doc, "7.3 Providências Pós-Plenário")
    pos = [
        "☐ Em caso de absolvição: alvará de soltura, comunicação às unidades, orientação",
        "☐ Em caso de condenação: interpor apelação (5 dias), requerer detração, regime",
        "☐ Em caso de condenação com pena ≤ tempo preso: requerer extinção da pena",
        "☐ Comunicar resultado aos familiares",
        "☐ Atualizar sistema OMBUDS com o resultado",
    ]
    for item in pos:
        add_body(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    # Rodapé final
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = p_final.add_run(f"Dossiê gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}\nDefensoria Pública do Estado da Bahia — 9ª Defensoria Pública\n7ª Regional — Camaçari — Bahia")
    run_final.font.size = Pt(9)
    run_final.font.color.rgb = RGBColor.from_string(MUTED)
    run_final.font.name = 'Verdana'

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"✅ Dossiê salvo em: {OUTPUT_FILE}")
    print(f"   Tamanho: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")


if __name__ == "__main__":
    generate_dossie()
