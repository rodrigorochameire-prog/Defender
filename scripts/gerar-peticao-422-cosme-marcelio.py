#!/usr/bin/env python3
"""
Petição de Diligências Preparatórias — Art. 422 CPP
Cosme Miguel Gomes de Souza e Marcélio Miguel de Souza
Processo: 8006374-21.2024.8.05.0039

Padrão DPE-BA — 7ª Regional de Camaçari
Gerado automaticamente pelo OMBUDS / Claude Code
"""

import os
from docx import Document
from docx.shared import Pt, Twips, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import numpy as np

# ============================================================
# PATHS
# ============================================================
LOGO_ORIG_PATH = "/Users/rodrigorochameire/Projetos/Defender/.claude/skills-cowork/dpe-ba-pecas/assets/dpe_logo.png"
LOGO_FADED_PATH = "/tmp/dpe_logo_faded_422.png"
OUTPUT_DIR = "/Users/rodrigorochameire/Meu Drive/1 - Defensoria 9ª DP/Protocolar"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Diligencias 422 - Cosme e Marcelio.docx")

# ============================================================
# PRÉ-PROCESSAR LOGO (opacidade 60%)
# ============================================================
img = Image.open(LOGO_ORIG_PATH).convert("RGBA")
arr = np.array(img, dtype=np.float64)
opacity = 0.60
white = np.full_like(arr[:, :, :3], 255.0)
arr[:, :, :3] = arr[:, :, :3] * opacity + white * (1 - opacity)
arr[:, :, 3] = 255
result = Image.fromarray(arr.astype(np.uint8)).convert("RGB")
result.save(LOGO_FADED_PATH)
print(f"Logo com opacidade {int(opacity * 100)}% gerada.")

# ============================================================
# CRIAR DOCUMENTO
# ============================================================
doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
font = style.font
font.name = 'Garamond'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- Configurar seção (margens) ---
section = doc.sections[0]
section.page_width = Twips(11906)
section.page_height = Twips(16838)
section.top_margin = Twips(2552)
section.bottom_margin = Twips(1134)
section.left_margin = Twips(1418)
section.right_margin = Twips(1134)
section.header_distance = Twips(567)
section.footer_distance = Twips(567)

# --- HEADER: logo ---
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hrun = hp.add_run()
hrun.add_picture(LOGO_FADED_PATH, width=Inches(1.777), height=Inches(1.101))

# --- FOOTER: rodapé institucional ---
footer = section.footer
footer.is_linked_to_previous = False

fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_after = Pt(0)
fp.paragraph_format.space_before = Pt(0)

# Borda superior do rodapé
pPr = fp._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
topBdr = OxmlElement('w:top')
topBdr.set(qn('w:val'), 'single')
topBdr.set(qn('w:sz'), '4')
topBdr.set(qn('w:space'), '1')
topBdr.set(qn('w:color'), '000000')
pBdr.append(topBdr)
pPr.append(pBdr)

spacing = pPr.find(qn('w:spacing'))
if spacing is None:
    spacing = OxmlElement('w:spacing')
    pPr.append(spacing)
spacing.set(qn('w:line'), '240')
spacing.set(qn('w:lineRule'), 'auto')

frun1 = fp.add_run("Defensoria Pública do Estado da Bahia")
frun1.font.name = 'Arial Narrow'
frun1.font.size = Pt(8)

fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp2.paragraph_format.space_after = Pt(0)
fp2.paragraph_format.space_before = Pt(0)
pPr2 = fp2._p.get_or_add_pPr()
spacing2 = OxmlElement('w:spacing')
spacing2.set(qn('w:line'), '240')
spacing2.set(qn('w:lineRule'), 'auto')
pPr2.append(spacing2)

frun2 = fp2.add_run("7ª Regional da DPE – Camaçari – Bahia.")
frun2.font.name = 'Arial Narrow'
frun2.font.size = Pt(8)

# ============================================================
# HELPERS
# ============================================================

def add_bold_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_after=Pt(0), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return p


def add_body_paragraph(doc, space_after=Pt(10)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Twips(720)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_run(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return run


def add_title(doc, text):
    add_empty_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return p


def add_empty_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Twips(720)
    return p


def add_subtitle(doc, text):
    """Subtítulo de item (ex: 6.1, 6.2) — bold, sem recuo, 6pt after"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return p


# ============================================================
# CONTEÚDO DO DOCUMENTO
# ============================================================

# --- ENDEREÇAMENTO ---
add_bold_paragraph(
    doc,
    "EXCELENTÍSSIMO(A) SENHOR(A) JUIZ(A) DE DIREITO DA VARA DO JÚRI "
    "E EXECUÇÕES PENAIS DA COMARCA DE CAMAÇARI — BAHIA",
    space_after=Pt(0)
)

# Duas linhas em branco
add_empty_line(doc)
add_empty_line(doc)

# --- EPÍGRAFE ---
add_bold_paragraph(doc, "Autos nº 8006374-21.2024.8.05.0039", space_after=Pt(20))

# Duas linhas em branco
add_empty_line(doc)
add_empty_line(doc)

# --- PREÂMBULO ---
p = add_body_paragraph(doc)
add_run(p, "COSME MIGUEL GOMES DE SOUZA", bold=True)
add_run(p, " e ")
add_run(p, "MARCÉLIO MIGUEL DE SOUZA", bold=True)
add_run(
    p,
    ", já qualificados nos autos em epígrafe, representados pela "
    "Defensoria Pública do Estado da Bahia, com fundamento no art. 134 "
    "da Constituição da República, por meio do defensor público subscritor, "
    "vêm respeitosamente perante Vossa Excelência, com fundamento no "
    "art. 422 do Código de Processo Penal, apresentar "
)
add_run(
    p,
    "ROL DE TESTEMUNHAS PARA PLENÁRIO E REQUERIMENTO DE DILIGÊNCIAS",
    bold=True
)
add_run(p, ", nos termos e pelas razões que seguem.")

# ============================================================
# I – DA OPORTUNIDADE PROCESSUAL
# ============================================================
add_title(doc, "I – DA OPORTUNIDADE PROCESSUAL")

p = add_body_paragraph(doc)
add_run(
    p,
    "Pronunciados os defendidos, os presentes autos foram remetidos ao "
    "Tribunal do Júri para julgamento em plenário. Nos termos do art. 422 "
    "do Código de Processo Penal:"
)

# Citação legal em itálico com recuo
p = add_body_paragraph(doc, space_after=Pt(10))
add_run(
    p,
    "\"Ao receber os autos, o presidente do Tribunal do Júri determinará "
    "a intimação do órgão do Ministério Público ou do querelante, no caso "
    "de queixa, e do defensor, para, no prazo de 5 (cinco) dias, "
    "apresentarem rol de testemunhas que irão depor em plenário, até o "
    "máximo de 5 (cinco), oportunidade em que poderão juntar documentos "
    "e requerer diligência.\"",
    italic=True
)

p = add_body_paragraph(doc)
add_run(
    p,
    "Trata-se, portanto, do momento processual adequado para que a Defesa "
    "apresente seu rol de testemunhas e requeira as diligências necessárias "
    "à instrução plenária, em atenção ao princípio da plenitude de defesa "
    "(art. 5º, XXXVIII, \"a\", da Constituição da República), que no "
    "Tribunal do Júri assume dimensão mais ampla do que a simples ampla "
    "defesa prevista no art. 5º, LV, da Carta Magna."
)

# ============================================================
# II – ROL DE TESTEMUNHAS PARA PLENÁRIO
# ============================================================
add_title(doc, "II – ROL DE TESTEMUNHAS PARA PLENÁRIO")

p = add_body_paragraph(doc)
add_run(
    p,
    "Nos termos do art. 422 do CPP, a Defesa apresenta o seguinte rol de "
    "testemunhas, comuns a ambos os defendidos:"
)

testemunhas = [
    ("ELISANGELA MARIA GOMES DE SOUZA", "irmã dos defendidos, ex-companheira do falecido",
     "Condomínio Pardais III, Bloco 82, Rua M, 101, Jardim Limoeiro, Camaçari-BA"),
    ("DORACI MARIANO FEITOSA", "vizinha, moradora do andar inferior",
     "Condomínio Pardais III, Jardim Limoeiro, Camaçari-BA"),
    ("SANDELEN MENDES DOS SANTOS", "vizinha do condomínio",
     "Condomínio Pardais III, Jardim Limoeiro, Camaçari-BA"),
]

add_empty_line(doc)
for i, (nome, qualif, endereco) in enumerate(testemunhas, 1):
    p = add_body_paragraph(doc, space_after=Pt(4))
    add_run(p, f"{i}. {nome}", bold=True)
    add_run(p, f" — {qualif}. {endereco}.")

add_empty_line(doc)
p = add_body_paragraph(doc)
add_run(
    p,
    "A Defesa reserva-se o direito de complementar o rol até o limite "
    "legal de 5 (cinco) testemunhas por defendido, após a conclusão de "
    "investigação defensiva em curso, oportunidade em que indicará a "
    "qualificação e o endereço das testemunhas remanescentes."
)

# ============================================================
# III – REQUERIMENTO DE DILIGÊNCIAS
# ============================================================
add_title(doc, "III – REQUERIMENTO DE DILIGÊNCIAS")

p = add_body_paragraph(doc)
add_run(
    p,
    "Em atenção à plenitude de defesa (art. 5º, XXXVIII, \"a\", CF) e "
    "aos poderes instrutórios do juiz presidente (art. 497, V, CPP), "
    "requer a Defesa as seguintes diligências:"
)

# 3.1
add_empty_line(doc)
p = add_body_paragraph(doc)
add_run(p, "3.1. ", bold=True)
add_run(p, "Requisição da ")
add_run(p, "folha de antecedentes criminais", bold=True)
add_run(
    p,
    " de Cláudio de Jesus Silva junto ao Instituto de Identificação "
    "Pedro Mello e à Polícia Civil do Estado da Bahia, a fim de documentar "
    "o histórico de violência do falecido, central à tese de legítima defesa."
)

# 3.2
p = add_body_paragraph(doc)
add_run(p, "3.2. ", bold=True)
add_run(p, "Requisição de ")
add_run(p, "todos os Boletins de Ocorrência", bold=True)
add_run(
    p,
    " registrados em nome de Cláudio de Jesus Silva como autor de violência "
    "doméstica, lesão corporal ou ameaça, e de Elisangela Maria Gomes de "
    "Souza como vítima, junto à Delegacia de Camaçari e à DEAM."
)

# 3.3
p = add_body_paragraph(doc)
add_run(p, "3.3. ", bold=True)
add_run(p, "Expedição de ofício à ")
add_run(p, "Vara de Violência Doméstica de Camaçari", bold=True)
add_run(
    p,
    " para informar se há ou houve medidas protetivas de urgência "
    "requeridas por Elisangela Maria Gomes de Souza em face de "
    "Cláudio de Jesus Silva."
)

# 3.4
p = add_body_paragraph(doc)
add_run(p, "3.4. ", bold=True)
add_run(p, "Determinação de ")
add_run(p, "laudo complementar ao exame necroscópico", bold=True)
add_run(
    p,
    " (IML Nina Rodrigues), para esclarecer: (a) se as lesões são "
    "compatíveis com reação defensiva; (b) a direção dos golpes e sua "
    "compatibilidade com a versão dos defendidos. Alternativamente, "
    "nomeação de assistente técnico da Defesa (art. 159, § 3º, CPP)."
)

# 3.5
p = add_body_paragraph(doc)
add_run(p, "3.5. ", bold=True)
add_run(p, "Requisição de ")
add_run(p, "registros médicos", bold=True)
add_run(
    p,
    " de Marcélio Miguel de Souza relativos à lesão no braço sofrida "
    "em 28/07/2018, junto a hospitais e UPAs de Camaçari — prova "
    "objetiva de que o falecido empregou arma branca contra os defendidos."
)


# ============================================================
# IV – PROVIDÊNCIAS PARA A SESSÃO PLENÁRIA
# ============================================================
add_title(doc, "IV – PROVIDÊNCIAS PARA A SESSÃO PLENÁRIA")

add_empty_line(doc)
p = add_body_paragraph(doc)
add_run(p, "4.1. ", bold=True)
add_run(
    p,
    "Apresentação dos defendidos em "
)
add_run(p, "roupa civil e sem algemas", bold=True)
add_run(
    p,
    ", nos termos da Súmula Vinculante nº 11/STF e art. 478, § 2º, CPP."
)

p = add_body_paragraph(doc)
add_run(p, "4.2. ", bold=True)
add_run(p, "Transferência antecipada", bold=True)
add_run(
    p,
    " de ambos os defendidos para estabelecimento prisional da Comarca de "
    "Camaçari, com antecedência mínima de 48 horas antes da sessão. "
    "Cosme encontra-se no Conjunto Penal de Juazeiro (a mais de 500 km) "
    "e Marcélio na Colônia Penal de Simões Filho, o que inviabiliza o "
    "atendimento prévio pela Defensoria sem a devida antecipação."
)

p = add_body_paragraph(doc)
add_run(p, "4.3. ", bold=True)
add_run(p, "Entrevista prévia reservada", bold=True)
add_run(
    p,
    " entre os defendidos e o defensor público, em local e horário "
    "adequados antes da sessão (art. 185, § 5º, CPP; art. 128, XI, "
    "LC nº 80/1994)."
)

# ============================================================
# V – DOS PEDIDOS
# ============================================================
add_title(doc, "V – DOS PEDIDOS")

p = add_body_paragraph(doc)
add_run(
    p,
    "Ante o exposto, a Defesa requer a Vossa Excelência:"
)

# Lista de pedidos
pedidos = [
    "O recebimento do rol de testemunhas apresentado no item II, com a intimação para comparecimento à sessão plenária;",
    "A concessão de prazo para complementação do rol, após a conclusão de investigação defensiva em curso;",
    "A requisição da folha de antecedentes criminais de Cláudio de Jesus Silva (item 3.1);",
    "A requisição de Boletins de Ocorrência relativos a Cláudio de Jesus Silva como autor e a Elisangela Maria Gomes de Souza como vítima (item 3.2);",
    "A expedição de ofício à Vara de Violência Doméstica de Camaçari sobre medidas protetivas (item 3.3);",
    "A determinação de laudo complementar ao IML ou nomeação de assistente técnico da Defesa (item 3.4);",
    "A requisição de registros médicos de Marcélio Miguel de Souza referentes à lesão sofrida em 28/07/2018 (item 3.5);",
    "A apresentação dos defendidos em roupa civil e sem algemas (item 4.1);",
    "A transferência antecipada dos defendidos para Camaçari, com mínimo de 48 horas de antecedência (item 4.2);",
    "A garantia de entrevista prévia reservada com o defensor público (item 4.3).",
]

for i, pedido in enumerate(pedidos, 1):
    p = add_body_paragraph(doc)
    add_run(p, f"{i}) ", bold=True)
    add_run(p, pedido)

# ============================================================
# FECHO, DATA E ASSINATURA
# ============================================================
add_empty_line(doc)

p = add_body_paragraph(doc)
add_run(p, "Nesses termos, pede deferimento.")

add_empty_line(doc)

p = add_body_paragraph(doc)
add_run(p, "Camaçari – BA, 06 de abril de 2026.")

add_empty_line(doc)
add_empty_line(doc)
add_empty_line(doc)

# Assinatura — centralizada, bold, espaçamento simples
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.space_before = Pt(0)
pPr_sig = p._p.get_or_add_pPr()
spacing_sig = OxmlElement('w:spacing')
spacing_sig.set(qn('w:line'), '240')
spacing_sig.set(qn('w:lineRule'), 'auto')
pPr_sig.append(spacing_sig)
run = p.add_run("Rodrigo Rocha Meire")
run.bold = True
run.font.name = 'Verdana'
run.font.size = Pt(12)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.space_before = Pt(0)
pPr_sig2 = p2._p.get_or_add_pPr()
spacing_sig2 = OxmlElement('w:spacing')
spacing_sig2.set(qn('w:line'), '240')
spacing_sig2.set(qn('w:lineRule'), 'auto')
pPr_sig2.append(spacing_sig2)
run2 = p2.add_run("Defensor Público")
run2.bold = True
run2.font.name = 'Verdana'
run2.font.size = Pt(12)

# ============================================================
# SALVAR
# ============================================================
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"\nPetição gerada com sucesso!")
print(f"Arquivo: {OUTPUT_FILE}")
