#!/usr/bin/env python3
"""
BIBLIOTECA VISUAL v3 — DPE-BA  |  7ª Regional – Camaçari
====================================================================
Helpers reutilizáveis para geração de documentos .docx profissionais.
Paleta soft, seções com acento lateral, timeline leve, teses em cards,
dashboard matplotlib e tabelas de análise estratégica.

COMO USAR:
  1. Copie este arquivo para o diretório de trabalho da sessão.
  2. Ajuste OUTPUT_PATH, LOGO_ORIG_PATH e LOGO_PATH.
  3. Execute doc = setup_document() para criar o documento base.
  4. Use os helpers abaixo para montar o conteúdo.
  5. Chame doc.save(OUTPUT_PATH) ao final.

HELPERS DISPONÍVEIS:
  Texto:
    add_title(doc, text)
    add_subtitle(doc, text)
    add_heading(doc, text)            ← faixa soft com borda esquerda navy
    add_subheading(doc, text)
    add_para(doc, text)
    add_bullet(doc, text, bold_label)
    add_mixed(doc, parts)             ← partes [(text, bold), ...]
    add_quote(doc, text)
    add_separator(doc)

  Tabelas visuais:
    build_hearing_banner(doc, processo, data_audiencia, vara, modo, link)
    build_painel_depoentes(doc, depoentes, reu_data, obs_operacional)
    build_timeline_phased(doc, fases)          ← timeline leve, 2 colunas
    build_comparison_table_colored(doc, headers, data)
    build_inconsistencias_table(doc, items)
    build_teses_table(doc, teses)              ← cards com borda colorida

  Gráficos matplotlib:
    generate_case_dashboard(chart_path, categorias, pontos_defesa, pontos_acusacao,
                            labels_pie, sizes_pie)

  Peças jurídicas (cabeçalho petição):
    build_cabecalho_peca(doc, processo, tipo_acao, vara, partes)
    build_requerimento_final(doc, texto_pedido)
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Twips, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image


# ════════════════════════════════════════════════════════════════════════════════
# PALETA PROFISSIONAL SOFT
# ════════════════════════════════════════════════════════════════════════════════

C_NAVY      = '1D3461'   # Azul corporativo — títulos, bordas
C_STEEL     = '2E6DA4'   # Azul médio — acentos
C_STEEL_LT  = 'D6E8F7'   # Azul muito suave — fundo de seção
C_WHITE     = 'FFFFFF'
C_OFF_WHITE = 'F8FAFB'   # Fundo de tabelas alternadas

# Banner de audiência
BAN_PROC_BG = 'EEF4FA'   # Fundo coluna processo  (azul mínimo)
BAN_DATA_BG = 'FDF4F4'   # Fundo coluna data      (rose mínimo)
BAN_VARA_BG = 'F3F0F8'   # Fundo coluna vara      (lavanda mínima)

# Status de intimações
STATUS_STYLES = {
    'ok':        ('D1FAE5', '065F46', '✅  INTIMADO(A)'),
    'curso':     ('FEF9C3', '713F12', '⚠️  DILIGÊNCIA EM CURSO'),
    'frustrada': ('FFE4E6', '881337', '❌  DILIGÊNCIA FRUSTRADA'),
    'nenhuma':   ('FEE2E2', '7F1D1D', '🔴  SEM INTIMAÇÃO'),
    'concluido': ('E5E7EB', '374151', '➖  FASE ENCERRADA'),
    'reu':       ('DBEAFE', '1E3A8A', ''),
}

# Fases da timeline
FASE_CORES = {
    'investigacao': ('DBEAFE', '1E3A8A'),
    'ip':           ('EDE9FE', '4C1D95'),
    'denuncia':     ('FEF9C3', '713F12'),
    'defesa':       ('D1FAE5', '14532D'),
    'audiencia':    ('FEE2E2', '7F1D1D'),
    'destaque':     ('1D3461', 'FFFFFF'),
    'neutro':       ('F8FAFC', '374151'),
    # Fases adicionais (VVD / Júri / Execução)
    'medida_protetiva': ('FCE7F3', '831843'),
    'execucao':     ('FEF3C7', '78350F'),
    'juri':         ('F0FDF4', '14532D'),
    'recurso':      ('FFF7ED', '7C2D12'),
}


# ════════════════════════════════════════════════════════════════════════════════
# SETUP DO DOCUMENTO (header, footer, margens)
# ════════════════════════════════════════════════════════════════════════════════

def preprocess_logo(logo_orig_path, logo_out_path, opacity=0.55):
    """Processa logo com opacidade sobre fundo branco."""
    img = Image.open(logo_orig_path).convert("RGBA")
    arr = np.array(img, dtype=np.float64)
    white = np.full_like(arr[:,:,:3], 255.0)
    arr[:,:,:3] = arr[:,:,:3] * opacity + white * (1 - opacity)
    arr[:,:,3] = 255
    Image.fromarray(arr.astype(np.uint8)).convert("RGB").save(logo_out_path)

def setup_document(logo_path):
    """Cria documento com estilos globais DPE-BA, header e footer."""
    doc = Document()
    # Estilo global
    style = doc.styles['Normal']
    style.font.name  = 'Verdana'
    style.font.size  = Pt(12)
    style.paragraph_format.line_spacing  = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Margens A4
    sec = doc.sections[0]
    sec.page_width    = Twips(11906); sec.page_height   = Twips(16838)
    sec.top_margin    = Twips(2552);  sec.bottom_margin = Twips(1134)
    sec.left_margin   = Twips(1418);  sec.right_margin  = Twips(1134)
    sec.header_distance = Twips(567); sec.footer_distance = Twips(567)

    # Header: logo
    header = sec.header; header.is_linked_to_previous = False
    hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run().add_picture(logo_path, width=Inches(1.777), height=Inches(1.101))

    # Footer: institucional
    footer = sec.footer; footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0); fp.paragraph_format.space_before = Pt(0)
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    topBdr = OxmlElement('w:top')
    topBdr.set(qn('w:val'), 'single'); topBdr.set(qn('w:sz'), '4')
    topBdr.set(qn('w:space'), '1');    topBdr.set(qn('w:color'), '8FA8C8')
    pBdr.append(topBdr); pPr.append(pBdr)
    r1 = fp.add_run("Defensoria Pública do Estado da Bahia")
    r1.font.name = 'Arial Narrow'; r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x4A,0x6F,0x8A)
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp2.paragraph_format.space_after = Pt(0); fp2.paragraph_format.space_before = Pt(0)
    r2 = fp2.add_run("7ª Regional da DPE – Camaçari – Bahia.")
    r2.font.name = 'Arial Narrow'; r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x4A,0x6F,0x8A)
    return doc


# ════════════════════════════════════════════════════════════════════════════════
# HELPERS DE CÉLULA (XML)
# ════════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:shd')): tcPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, font_size=9, color_hex=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ''; p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = 'Garamond'; r.font.size = Pt(font_size)
    if color_hex:
        r.font.color.rgb = RGBColor(
            int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))

def set_col_width(cell, width_twips):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips)); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_cell_border_left(cell, color_hex, sz=24):
    """Borda esquerda colorida — acento de fase/seção."""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcB = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), str(sz))
    left.set(qn('w:color'), color_hex); left.set(qn('w:space'), '0')
    tcB.append(left)
    for side in ('top','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tcB.append(el)
    tcPr.append(tcB)

def set_cell_border_top(cell, color_hex, sz=8):
    """Borda superior colorida."""
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:tcBorders')): tcPr.remove(e)
    tcB = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), str(sz))
    top.set(qn('w:color'), color_hex); top.set(qn('w:space'), '0')
    tcB.append(top)
    for side in ('bottom','left','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'none'); tcB.append(el)
    tcPr.append(tcB)


# ════════════════════════════════════════════════════════════════════════════════
# HELPERS DE TEXTO
# ════════════════════════════════════════════════════════════════════════════════

def add_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.name = 'Garamond'; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = 'Garamond'; r.font.size = Pt(10)
    r.italic = True; r.font.color.rgb = RGBColor(0x4A,0x6F,0x8A)
    return p

def add_heading(doc, text, bg_color=None, border_color=None):
    """Faixa soft com borda esquerda navy — substitui heading sublinhado."""
    bg = bg_color or C_STEEL_LT
    bc = border_color or C_NAVY
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg); set_cell_border_left(cell, bc, sz=28)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Twips(220)
    r = p.add_run(text); r.bold = True; r.font.name = 'Garamond'; r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    doc.add_paragraph()
    return tbl

def add_subheading(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.name = 'Garamond'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    return p

def add_para(doc, text, first_line_indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if first_line_indent: p.paragraph_format.first_line_indent = Twips(720)
    r = p.add_run(text); r.font.name = 'Garamond'; r.font.size = Pt(12)
    return p

def add_bullet(doc, text, bold_label=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Twips(360)
    if bold_label:
        r1 = p.add_run(f"{bold_label} "); r1.bold = True
        r1.font.name = 'Garamond'; r1.font.size = Pt(12)
        r2 = p.add_run(text); r2.font.name = 'Garamond'; r2.font.size = Pt(12)
    else:
        r = p.add_run(f"• {text}"); r.font.name = 'Garamond'; r.font.size = Pt(12)
    return p

def add_mixed(doc, parts, first_line_indent=True):
    """parts: list of (text, bold). Ex: [('Label: ', True), ('corpo', False)]"""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if first_line_indent: p.paragraph_format.first_line_indent = Twips(720)
    for text, bold in parts:
        r = p.add_run(text); r.bold = bold
        r.font.name = 'Garamond'; r.font.size = Pt(12)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Twips(720); p.paragraph_format.right_indent = Twips(720)
    r = p.add_run(f'"{text}"'); r.italic = True
    r.font.name = 'Garamond'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x4A,0x5C,0x6A)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run("─" * 70); r.font.name = 'Garamond'; r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xB0,0xC4,0xDE)
    return p


# ════════════════════════════════════════════════════════════════════════════════
# DASHBOARD MATPLOTLIB
# ════════════════════════════════════════════════════════════════════════════════

def generate_case_dashboard(chart_path,
                             categorias=None, pontos_defesa=None, pontos_acusacao=None,
                             labels_pie=None, sizes_pie=None,
                             titulo_chart="Avaliação do Caso"):
    """
    Gera PNG dual: balanço probatório (barras) + status de intimações (pizza).

    Parâmetros:
      chart_path       — caminho de saída do PNG
      categorias       — lista de categorias probatórias (eixo X do gráfico de barras)
      pontos_defesa    — lista de floats 0-1 (força da defesa em cada categoria)
      pontos_acusacao  — lista de floats 0-1 (força da acusação em cada categoria)
      labels_pie       — rótulos do gráfico de pizza (status de depoentes)
      sizes_pie        — contagens para o pizza
      titulo_chart     — título geral do dashboard

    Exemplo VVD:
      generate_case_dashboard(
          chart_path='/sessions/.../dashboard.png',
          categorias=['Materialidade','Provas\ndocumentais','Relato\nvítima','Versão\nréu','Contexto\nVD'],
          pontos_defesa   =[0.70, 0.60, 0.30, 0.80, 0.50],
          pontos_acusacao =[0.40, 0.50, 0.85, 0.25, 0.70],
          labels_pie      =['Intimados','Em curso','Frustrada','Sem intimação'],
          sizes_pie       =[2, 1, 1, 1],
          titulo_chart    ='Avaliação do Caso VVD',
      )
    """
    if categorias is None:
        categorias = ['Materialidade\n(laudo)', 'Prova técnica', 'Testemunhas\ndiretas', 'Credib.\nofendida', 'Versão\ndefesa']
    if pontos_defesa is None:
        pontos_defesa   = [0.70, 0.65, 0.55, 0.40, 0.75]
    if pontos_acusacao is None:
        pontos_acusacao = [0.35, 0.30, 0.55, 0.75, 0.30]
    if labels_pie is None:
        labels_pie = ['Intimados\n(prontos)', 'Diligência\nem curso', 'Diligência\nfrustrada', 'Sem\nintimação']
    if sizes_pie is None:
        sizes_pie = [2, 2, 1, 1]

    plt.rcParams.update({'font.family': 'DejaVu Sans', 'font.size': 8,
                         'axes.spines.top': False, 'axes.spines.right': False})

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 3.6))
    fig.patch.set_facecolor('#F8FAFB')

    # Barras
    x = np.arange(len(categorias)); w = 0.35
    bars_d = ax1.bar(x - w/2, pontos_defesa,   w, label='Defesa',   color='#4A90C4', alpha=0.85, zorder=3)
    bars_a = ax1.bar(x + w/2, pontos_acusacao, w, label='Acusação', color='#E07070', alpha=0.85, zorder=3)
    ax1.set_xticks(x); ax1.set_xticklabels(categorias, fontsize=7.5)
    ax1.set_ylim(0, 1.05); ax1.set_yticks([0, 0.25, 0.50, 0.75, 1.0])
    ax1.set_yticklabels(['0%','25%','50%','75%','100%'], fontsize=7)
    ax1.yaxis.grid(True, linestyle='--', alpha=0.4, zorder=0)
    ax1.set_facecolor('#FFFFFF')
    ax1.set_title('Avaliação Probatória — Defesa × Acusação',
                  fontsize=9, fontweight='bold', color='#1D3461', pad=8)
    ax1.legend(fontsize=7.5, loc='upper right')
    ax1.axhline(0.5, color='#CBD5E0', linewidth=0.8, linestyle=':')
    for bar in bars_d:
        h = bar.get_height()
        ax1.text(bar.get_x()+bar.get_width()/2., h+0.01, f'{h:.0%}', ha='center', va='bottom', fontsize=6.5, color='#2C5282')
    for bar in bars_a:
        h = bar.get_height()
        ax1.text(bar.get_x()+bar.get_width()/2., h+0.01, f'{h:.0%}', ha='center', va='bottom', fontsize=6.5, color='#9B2335')

    # Pizza
    colors_pie = ['#68D391','#F6C90E','#F87171','#FCA5A5'][:len(sizes_pie)]
    explode = [0.03 if i == 0 else 0.06 for i in range(len(sizes_pie))]
    wedges, texts, autotexts = ax2.pie(
        sizes_pie, labels=labels_pie, colors=colors_pie, explode=explode,
        autopct='%1.0f%%', startangle=90,
        textprops={'fontsize': 7.5}, pctdistance=0.72,
        wedgeprops={'linewidth': 0.8, 'edgecolor': 'white'})
    for at in autotexts: at.set_fontsize(7); at.set_color('#2D3748'); at.set_fontweight('bold')
    ax2.set_title('Status de Intimações', fontsize=9, fontweight='bold', color='#1D3461', pad=8)
    ax2.set_facecolor('#FFFFFF')

    fig.text(0.5, -0.02, titulo_chart, ha='center', fontsize=7.5, color='#718096', style='italic')
    plt.tight_layout(pad=1.5)
    plt.savefig(chart_path, dpi=150, bbox_inches='tight', facecolor='#F8FAFB', edgecolor='none')
    plt.close()
    return chart_path


# ════════════════════════════════════════════════════════════════════════════════
# BANNER DE AUDIÊNCIA
# ════════════════════════════════════════════════════════════════════════════════

def build_hearing_banner(doc, processo, data_audiencia, vara, modo, link=None):
    """
    Banner 3 colunas — backgrounds suaves, borda superior colorida por coluna.
    Uso: build_hearing_banner(doc, '8002212-...', '17/03/2026 · 11h00', 'Vara VD — Camaçari', 'HÍBRIDA', 'https://...')
    """
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Table Grid'

    c0 = tbl.rows[0].cells[0]; set_cell_bg(c0, BAN_PROC_BG); set_cell_border_top(c0, C_NAVY, sz=12)
    c0.text = ''; p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(8); p0.paragraph_format.space_after = Pt(8)
    r0a = p0.add_run("📁  PROCESSO\n"); r0a.font.name = 'Garamond'; r0a.font.size = Pt(7)
    r0a.font.color.rgb = RGBColor(0x4A,0x6F,0x8A)
    r0b = p0.add_run(processo); r0b.bold = True; r0b.font.name = 'Garamond'; r0b.font.size = Pt(9)
    r0b.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    set_col_width(c0, 2400)

    c1 = tbl.rows[0].cells[1]; set_cell_bg(c1, BAN_DATA_BG); set_cell_border_top(c1, 'C53030', sz=12)
    c1.text = ''; p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(6); p1.paragraph_format.space_after = Pt(2)
    r1a = p1.add_run("🔴  AUDIÊNCIA\n"); r1a.font.name = 'Garamond'; r1a.font.size = Pt(8)
    r1a.font.color.rgb = RGBColor(0xC5,0x30,0x30)
    r1b = p1.add_run(data_audiencia); r1b.bold = True; r1b.font.name = 'Garamond'; r1b.font.size = Pt(15)
    r1b.font.color.rgb = RGBColor(0xC5,0x30,0x30)
    set_col_width(c1, 2800)

    c2 = tbl.rows[0].cells[2]; set_cell_bg(c2, BAN_VARA_BG); set_cell_border_top(c2, C_STEEL, sz=12)
    c2.text = ''; p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8); p2.paragraph_format.space_after = Pt(2)
    r2a = p2.add_run(f"⚖️  {vara}\n"); r2a.font.name = 'Garamond'; r2a.font.size = Pt(8)
    r2a.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    r2b = p2.add_run(f"📡  {modo}"); r2b.bold = True; r2b.font.name = 'Garamond'; r2b.font.size = Pt(9)
    r2b.font.color.rgb = RGBColor(0x2E,0x6D,0xA4)
    if link:
        p2b = c2.add_paragraph(); p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2b.paragraph_format.space_before = Pt(2); p2b.paragraph_format.space_after = Pt(6)
        r2c = p2b.add_run(link); r2c.font.name = 'Garamond'; r2c.font.size = Pt(6.5)
        r2c.font.color.rgb = RGBColor(0x5B,0x8A,0xB5)
    set_col_width(c2, 2400)
    doc.add_paragraph()
    return tbl


# ════════════════════════════════════════════════════════════════════════════════
# PAINEL DE DEPOENTES
# ════════════════════════════════════════════════════════════════════════════════

def build_painel_depoentes(doc, depoentes, reu_data, obs_operacional):
    """
    Tabela de depoentes com status de intimação + tabela do réu + alerta operacional.

    depoentes: list of dict com chaves:
      name, papel, delegacia_sim (bool), delegacia_data, juizo_sim (bool), juizo_data,
      status ('ok'|'curso'|'frustrada'|'nenhuma'|'concluido'), status_detalhe, obs

    reu_data: dict com chaves:
      citado, intimado_audiencia, preso, endereco_obs

    obs_operacional: string com alertas para a audiência
    """
    tbl = doc.add_table(rows=1, cols=5); tbl.style = 'Table Grid'
    widths = [2100, 1200, 900, 900, 2500]
    col_labels = ['NOME', 'PAPEL', 'DELEGACIA', 'JUÍZO', 'STATUS — PRÓXIMA AUDIÊNCIA']
    for i, (label, w) in enumerate(zip(col_labels, widths)):
        set_cell_bg(tbl.rows[0].cells[i], C_NAVY)
        set_cell_text(tbl.rows[0].cells[i], label, bold=True, font_size=8.5, color_hex=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_col_width(tbl.rows[0].cells[i], w)

    for idx, d in enumerate(depoentes):
        row_bg = C_OFF_WHITE if idx % 2 == 0 else C_WHITE
        fundo, texto_cor, status_label = STATUS_STYLES.get(d['status'], STATUS_STYLES['nenhuma'])
        row = tbl.add_row().cells
        set_cell_bg(row[0], row_bg)
        set_cell_text(row[0], d['name'], bold=True, font_size=9, color_hex=C_NAVY)
        if d.get('obs'):
            p2 = row[0].add_paragraph(); p2.paragraph_format.space_before = Pt(1)
            r2 = p2.add_run(d['obs']); r2.font.name = 'Garamond'; r2.font.size = Pt(7.5)
            r2.italic = True; r2.font.color.rgb = RGBColor(0x71,0x80,0x96)
        set_col_width(row[0], widths[0])

        set_cell_bg(row[1], row_bg)
        set_cell_text(row[1], d['papel'], font_size=8.5, color_hex='374151')
        bg_del = 'D1FAE5' if d['delegacia_sim'] else 'FEE2E2'
        set_cell_bg(row[2], bg_del)
        set_cell_text(row[2], ('✔  '+d.get('delegacia_data','')) if d['delegacia_sim'] else '—',
                      font_size=8, color_hex='065F46' if d['delegacia_sim'] else '7F1D1D',
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        bg_jui = 'D1FAE5' if d['juizo_sim'] else 'FEF9C3'
        set_cell_bg(row[3], bg_jui)
        set_cell_text(row[3], ('✔  '+d.get('juizo_data','')) if d['juizo_sim'] else 'Pendente',
                      font_size=8, color_hex='065F46' if d['juizo_sim'] else '713F12',
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(row[4], fundo)
        set_cell_text(row[4], status_label, bold=True, font_size=9, color_hex=texto_cor)
        if d.get('status_detalhe'):
            p_det = row[4].add_paragraph(); p_det.paragraph_format.space_before = Pt(1)
            r_det = p_det.add_run(d['status_detalhe']); r_det.font.name = 'Garamond'; r_det.font.size = Pt(7.5)
            r_det.font.color.rgb = RGBColor(int(texto_cor[0:2],16),int(texto_cor[2:4],16),int(texto_cor[4:6],16))

    doc.add_paragraph()

    # Tabela do réu
    tbl_r = doc.add_table(rows=1, cols=2); tbl_r.style = 'Table Grid'
    hdr = tbl_r.rows[0].cells; set_cell_bg(hdr[0], C_NAVY)
    set_cell_text(hdr[0], f"👤  RÉU — {reu_data.get('nome','ASSISTIDO').upper()}",
                  bold=True, font_size=10, color_hex=C_WHITE)
    set_col_width(hdr[0], 3800); hdr[0].merge(hdr[1])
    campos = [('Citação', reu_data.get('citado','')),
              ('Intimação para a audiência', reu_data.get('intimado_audiencia','')),
              ('Status de prisão', reu_data.get('preso','')),
              ('Endereço / observações', reu_data.get('endereco_obs',''))]
    for i, (label, valor) in enumerate(campos):
        row = tbl_r.add_row().cells; row_bg = C_STEEL_LT if i%2==0 else C_WHITE
        set_cell_bg(row[0], row_bg); set_cell_text(row[0], label, bold=True, font_size=9, color_hex=C_NAVY)
        set_col_width(row[0], 1900); set_cell_bg(row[1], C_WHITE)
        set_cell_text(row[1], valor, font_size=9, color_hex='2D3748'); set_col_width(row[1], 5700)
    doc.add_paragraph()

    # Alerta operacional
    tbl_al = doc.add_table(rows=1, cols=1); tbl_al.style = 'Table Grid'
    c_al = tbl_al.rows[0].cells[0]; set_cell_bg(c_al, 'FFFBEB')
    set_cell_border_left(c_al, 'B45309', sz=24); c_al.text = ''
    p_al = c_al.paragraphs[0]; p_al.paragraph_format.space_before = Pt(6)
    p_al.paragraph_format.left_indent = Twips(180)
    rh = p_al.add_run("⚠  ALERTA OPERACIONAL\n"); rh.bold = True
    rh.font.name = 'Garamond'; rh.font.size = Pt(9)
    rh.font.color.rgb = RGBColor(0xB4,0x53,0x09)
    rb = p_al.add_run(obs_operacional); rb.font.name = 'Garamond'; rb.font.size = Pt(9)
    rb.font.color.rgb = RGBColor(0x44,0x33,0x10)


# ════════════════════════════════════════════════════════════════════════════════
# TIMELINE LEVE (v3 — sem eixo escuro)
# ════════════════════════════════════════════════════════════════════════════════

def build_timeline_phased(doc, fases):
    """
    Timeline leve, 2 colunas.
    Col 0 (1200 twips): data em negrito na cor da fase — fundo branco.
    Col 1 (6400 twips): ícone + descrição — fundo branco, borda esquerda fina da cor da fase.
    Cabeçalho de fase: fundo suave + título bold na cor escura (sem banda escura pesada).

    fases: list of (fase_key, titulo_fase, periodo, [(data, icone, descricao), ...])

    Fases disponíveis:
      'investigacao', 'ip', 'denuncia', 'defesa', 'audiencia', 'destaque', 'neutro',
      'medida_protetiva', 'execucao', 'juri', 'recurso'

    Exemplo VVD:
      build_timeline_phased(doc, [
          ('investigacao', '🔵  FASE 1 — FLAGRANTE / BO', 'Nov 2022', [
              ('27/11/2022', '🚔', 'Boletim de Ocorrência — violência doméstica'),
          ]),
          ('medida_protetiva', '🩷  FASE 2 — MEDIDAS PROTETIVAS', 'Dez 2022', [
              ('05/12/2022', '🛡️', 'Medida protetiva concedida — afastamento do lar'),
          ]),
          ('denuncia', '🟡  FASE 3 — DENÚNCIA', 'Mar 2023', [...]),
          ('destaque', '🔴  HOJE — AUDIÊNCIA', '17/03/2026', [...]),
      ])
    """
    COL_DATE = 1200; COL_DESC = 6400
    tbl = doc.add_table(rows=0, cols=2); tbl.style = 'Table Grid'

    for fase_key, titulo_fase, periodo, eventos in fases:
        fundo_claro, cor_escura = FASE_CORES.get(fase_key, FASE_CORES['neutro'])

        # Cabeçalho de fase
        hdr = tbl.add_row().cells
        set_cell_bg(hdr[0], fundo_claro); hdr[0].text = ''
        ph0 = hdr[0].paragraphs[0]; ph0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph0.paragraph_format.space_before = Pt(6); ph0.paragraph_format.space_after = Pt(6)
        rp = ph0.add_run(periodo); rp.bold = True; rp.italic = True
        rp.font.name = 'Garamond'; rp.font.size = Pt(8)
        rp.font.color.rgb = RGBColor(int(cor_escura[0:2],16),int(cor_escura[2:4],16),int(cor_escura[4:6],16))
        set_col_width(hdr[0], COL_DATE)

        set_cell_bg(hdr[1], fundo_claro); set_cell_border_left(hdr[1], cor_escura, sz=20); hdr[1].text = ''
        ph1 = hdr[1].paragraphs[0]; ph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ph1.paragraph_format.space_before = Pt(6); ph1.paragraph_format.space_after = Pt(6)
        ph1.paragraph_format.left_indent = Twips(160)
        rt = ph1.add_run(titulo_fase); rt.bold = True
        rt.font.name = 'Garamond'; rt.font.size = Pt(9.5)
        rt.font.color.rgb = RGBColor(int(cor_escura[0:2],16),int(cor_escura[2:4],16),int(cor_escura[4:6],16))
        set_col_width(hdr[1], COL_DESC)

        for data, icone, descricao in eventos:
            row = tbl.add_row().cells
            set_cell_bg(row[0], C_WHITE); row[0].text = ''
            pd = row[0].paragraphs[0]; pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pd.paragraph_format.space_before = Pt(5); pd.paragraph_format.space_after = Pt(5)
            rd = pd.add_run(data); rd.bold = True; rd.font.name = 'Garamond'; rd.font.size = Pt(8)
            rd.font.color.rgb = RGBColor(int(cor_escura[0:2],16),int(cor_escura[2:4],16),int(cor_escura[4:6],16))
            set_col_width(row[0], COL_DATE)

            set_cell_bg(row[1], C_WHITE); set_cell_border_left(row[1], cor_escura, sz=12); row[1].text = ''
            pdesc = row[1].paragraphs[0]; pdesc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pdesc.paragraph_format.space_before = Pt(5); pdesc.paragraph_format.space_after = Pt(5)
            pdesc.paragraph_format.left_indent = Twips(160)
            ri = pdesc.add_run(icone + "  "); ri.font.name = 'Segoe UI Emoji'; ri.font.size = Pt(10)
            rd2 = pdesc.add_run(descricao); rd2.font.name = 'Garamond'; rd2.font.size = Pt(8.5)
            rd2.font.color.rgb = RGBColor(0x2D,0x37,0x48)
            set_col_width(row[1], COL_DESC)

    return tbl


# ════════════════════════════════════════════════════════════════════════════════
# TABELA COMPARATIVA
# ════════════════════════════════════════════════════════════════════════════════

def build_comparison_table_colored(doc, headers, data):
    """
    Tabela de confronto de versões / pontos fáticos.
    headers: [col0_title, col1_title, col2_title]
    data: list of (label, col2, col3, tipo)
    tipo: 'divergencia' | 'convergencia' | 'neutro'
    """
    TIPO_CORES = {
        'divergencia':  ('FFF0F0', '882B2B', '⚔'),
        'convergencia': ('F0FFF4', '1A5C36', '✔'),
        'neutro':       ('F8FAFC', '374151', '—'),
    }
    tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Table Grid'
    col_widths = [2200, 2800, 2600]
    for i, (label, w) in enumerate(zip(headers, col_widths)):
        set_cell_bg(tbl.rows[0].cells[i], C_NAVY)
        set_cell_text(tbl.rows[0].cells[i], label, bold=True, font_size=8.5, color_hex=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_col_width(tbl.rows[0].cells[i], w)

    for i, row_data in enumerate(data):
        label, col2, col3 = row_data[:3]
        tipo = row_data[3] if len(row_data) > 3 else 'neutro'
        fundo, texto_cor, icon = TIPO_CORES.get(tipo, TIPO_CORES['neutro'])
        row_bg = C_OFF_WHITE if i%2==0 else C_WHITE
        row = tbl.add_row().cells

        set_cell_bg(row[0], fundo); row[0].text = ''
        p = row[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        ri = p.add_run(f"{icon}  "); ri.bold = True; ri.font.name = 'Garamond'; ri.font.size = Pt(9)
        ri.font.color.rgb = RGBColor(int(texto_cor[0:2],16),int(texto_cor[2:4],16),int(texto_cor[4:6],16))
        rl = p.add_run(label); rl.bold = True; rl.font.name = 'Garamond'; rl.font.size = Pt(8.5)
        rl.font.color.rgb = RGBColor(int(texto_cor[0:2],16),int(texto_cor[2:4],16),int(texto_cor[4:6],16))
        set_col_width(row[0], col_widths[0])
        set_cell_bg(row[1], row_bg); set_cell_text(row[1], col2, font_size=8.5, color_hex='2D3748')
        set_col_width(row[1], col_widths[1])
        set_cell_bg(row[2], row_bg); set_cell_text(row[2], col3, font_size=8.5, color_hex='2D3748')
        set_col_width(row[2], col_widths[2])
    return tbl


# ════════════════════════════════════════════════════════════════════════════════
# TABELA DE INCONSISTÊNCIAS / VULNERABILIDADES
# ════════════════════════════════════════════════════════════════════════════════

def build_inconsistencias_table(doc, items):
    """
    items: list of (numero, nivel, titulo, conteudo)
    nivel: 'critico' | 'alto' | 'medio' | 'info'

    Uso em VVD: vulnerabilidades da acusação, pontos para cross-examination.
    Uso em Júri: pontos críticos da pronúncia, argumentos para plenário.
    """
    NIVEL_STYLES = {
        'critico': ('7F1D1D', 'FFF0F0', '🔴 CRÍTICO'),
        'alto':    ('78350F', 'FFFBEB', '🟠 ALTO'),
        'medio':   ('14532D', 'F0FFF4', '🟡 MÉDIO'),
        'info':    ('1E3A8A', 'EEF4FF', 'ℹ️ OBS'),
    }
    tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Table Grid'
    cols_cfg = [('Nº', 400), ('NÍVEL', 1100), ('INCONSISTÊNCIA / PONTO ESTRATÉGICO', 6100)]
    for i, (label, w) in enumerate(cols_cfg):
        set_cell_bg(tbl.rows[0].cells[i], C_NAVY)
        set_cell_text(tbl.rows[0].cells[i], label, bold=True, font_size=8.5, color_hex=C_WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_col_width(tbl.rows[0].cells[i], w)

    for numero, nivel, titulo, conteudo in items:
        cor_txt, cor_fundo, nivel_label = NIVEL_STYLES.get(nivel, NIVEL_STYLES['info'])
        row = tbl.add_row().cells
        set_cell_bg(row[0], cor_fundo)
        set_cell_text(row[0], str(numero), bold=True, font_size=11, color_hex=cor_txt,
                      align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(row[0], 400)
        set_cell_bg(row[1], cor_fundo)
        set_cell_text(row[1], nivel_label, bold=True, font_size=8, color_hex=cor_txt,
                      align=WD_ALIGN_PARAGRAPH.CENTER); set_col_width(row[1], 1100)
        set_cell_bg(row[2], C_WHITE); set_cell_border_left(row[2], cor_txt, sz=16)
        row[2].text = ''; p = row[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Twips(80)
        r1 = p.add_run(titulo + ":  "); r1.bold = True; r1.font.name = 'Garamond'; r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(int(cor_txt[0:2],16),int(cor_txt[2:4],16),int(cor_txt[4:6],16))
        r2 = p.add_run(conteudo); r2.font.name = 'Garamond'; r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(0x2D,0x37,0x48); set_col_width(row[2], 6100)
    return tbl


# ════════════════════════════════════════════════════════════════════════════════
# TESES DEFENSIVAS — CARDS LEVES
# ════════════════════════════════════════════════════════════════════════════════

def build_teses_table(doc, teses):
    """
    Cards leves: borda esquerda colorida, número inline, sem bloco escuro.

    teses: list of (numero, titulo, conteudo)

    Uso adaptável:
      - audiência criminal: teses de mérito para instrução
      - VVD: fundamentos para alegações finais / resposta à acusação
      - Júri: argumentos para plenário / quesitos favoráveis
      - Criminal: teses para HC / revogação de prisão
    """
    CORES_TESE = [C_NAVY, '14532D', '78350F', '44337A', '7F1D1D', '065F46']

    for idx, (numero, titulo, conteudo) in enumerate(teses):
        cor = CORES_TESE[idx % len(CORES_TESE)]
        tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, C_OFF_WHITE); set_cell_border_left(cell, cor, sz=24); cell.text = ''

        ph = cell.paragraphs[0]; ph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ph.paragraph_format.space_before = Pt(6); ph.paragraph_format.space_after = Pt(2)
        ph.paragraph_format.left_indent = Twips(200)
        rn = ph.add_run(f"#{numero}  —  "); rn.bold = True; rn.font.name = 'Garamond'; rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor(int(cor[0:2],16),int(cor[2:4],16),int(cor[4:6],16))
        rt = ph.add_run(titulo); rt.bold = True; rt.font.name = 'Garamond'; rt.font.size = Pt(9.5)
        rt.font.color.rgb = RGBColor(int(cor[0:2],16),int(cor[2:4],16),int(cor[4:6],16))

        pb = cell.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pb.paragraph_format.space_before = Pt(2); pb.paragraph_format.space_after = Pt(7)
        pb.paragraph_format.left_indent = Twips(200)
        rb = pb.add_run(conteudo); rb.font.name = 'Garamond'; rb.font.size = Pt(9)
        rb.font.color.rgb = RGBColor(0x4A,0x55,0x68)
        doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════════
# CABEÇALHO DE PEÇA JURÍDICA
# ════════════════════════════════════════════════════════════════════════════════

def build_cabecalho_peca(doc, tipo_peca, processo, vara, partes):
    """
    Cabeçalho padrão de petição — útil para qualquer peça processual.

    tipo_peca : str  — ex: 'ALEGAÇÕES FINAIS', 'RESPOSTA À ACUSAÇÃO', 'HABEAS CORPUS'
    processo  : str  — número do processo
    vara      : str  — vara / juízo
    partes    : dict — {'reu': 'Nome Réu', 'vitima': 'Nome Vítima', 'mp': 'Promotor'}

    Exemplo:
      build_cabecalho_peca(doc,
          tipo_peca='ALEGAÇÕES FINAIS (MEMORIAIS)',
          processo='0000000-00.0000.8.05.0000',
          vara='Vara de Violência Doméstica — Camaçari/BA',
          partes={'reu': 'Nome do Réu', 'vitima': 'Nome da Vítima'})
    """
    add_title(doc, tipo_peca)
    doc.add_paragraph()
    # Box de identificação
    tbl = doc.add_table(rows=0, cols=2); tbl.style = 'Table Grid'
    campos = [('PROCESSO', processo), ('VARA / JUÍZO', vara)]
    campos += [(k.upper(), v) for k, v in partes.items()]
    for i, (k, v) in enumerate(campos):
        row = tbl.add_row().cells
        bg = C_STEEL_LT if i%2==0 else C_WHITE
        set_cell_bg(row[0], bg); set_cell_text(row[0], k, bold=True, font_size=9, color_hex=C_NAVY)
        set_col_width(row[0], 2000)
        set_cell_bg(row[1], C_WHITE); set_cell_text(row[1], v, font_size=9, color_hex='2D3748')
        set_col_width(row[1], 5600)
    doc.add_paragraph()


def build_requerimento_final(doc, texto_pedido, local_data="Camaçari, ____ de ________ de 2026.",
                              defensor="Rodrigo Rocha Meire", cargo="Defensor Público"):
    """
    Bloco final de requerimento: 'Ante o exposto, requer...' + assinatura.
    """
    doc.add_paragraph()
    add_separator(doc)
    add_subheading(doc, "DOS PEDIDOS")
    add_para(doc, texto_pedido)
    doc.add_paragraph()
    p_dt = doc.add_paragraph(); p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dt.add_run(local_data).font.name = 'Garamond'
    doc.add_paragraph()
    p_as = doc.add_paragraph(); p_as.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_as = p_as.add_run(defensor); r_as.bold = True
    r_as.font.name = 'Garamond'; r_as.font.size = Pt(12)
    r_as.font.color.rgb = RGBColor(0x1D,0x34,0x61)
    p_cr = doc.add_paragraph(); p_cr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cr.add_run(cargo).font.name = 'Garamond'


# ════════════════════════════════════════════════════════════════════════════════
# EXEMPLO DE USO — substituir pela lógica real do caso
# ════════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    # Ajuste estes caminhos para a sessão atual
    SESSION = "/sessions/NOME-DA-SESSAO"
    LOGO_ORIG = f"{SESSION}/mnt/.skills/skills/analise-audiencias/assets/dpe_logo.png"
    LOGO_OUT  = f"{SESSION}/dpe_logo_faded.png"
    OUTPUT    = f"{SESSION}/mnt/PASTA-DO-CASO/documento.docx"
    CHART_OUT = f"{SESSION}/case_dashboard.png"

    preprocess_logo(LOGO_ORIG, LOGO_OUT)
    doc = setup_document(LOGO_OUT)

    # Dashboard (opcional — para relatórios de audiência)
    generate_case_dashboard(
        chart_path=CHART_OUT,
        categorias=['Materialidade', 'Prova técnica', 'Testemunhas', 'Relato\nvítima', 'Versão\nréu'],
        pontos_defesa   =[0.85, 0.80, 0.60, 0.40, 0.75],
        pontos_acusacao =[0.25, 0.20, 0.50, 0.75, 0.30],
        labels_pie=['Intimados', 'Em curso', 'Frustrada', 'Sem intim.'],
        sizes_pie=[3, 2, 1, 1],
        titulo_chart='Avaliação — Caso Exemplo'
    )

    # Título + banner
    add_title(doc, "RELATÓRIO PARA AUDIÊNCIA — EXEMPLO")
    build_hearing_banner(doc, 'PROCESSO-NUMERO', 'DD/MM/AAAA · HHhMM',
                         'Vara VD — Camaçari/BA', 'PRESENCIAL')

    # Dashboard embutido
    add_heading(doc, "AVALIAÇÃO DO CASO")
    p_ch = doc.add_paragraph(); p_ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ch.add_run().add_picture(CHART_OUT, width=Inches(6.2))

    # Cronologia
    add_heading(doc, "CRONOLOGIA PROCESSUAL")
    build_timeline_phased(doc, [
        ('investigacao', '🔵  FASE 1 — FLAGRANTE', 'Data', [
            ('DD/MM/AAAA', '🚔', 'Descrição do ato processual'),
        ]),
        ('destaque', '🔴  HOJE — AUDIÊNCIA', 'DD/MM/AAAA', [
            ('DD/MM/AAAA', '🔴', 'AUDIÊNCIA DE INSTRUÇÃO — HHhMM'),
        ]),
    ])

    # Inconsistências
    add_heading(doc, "INCONSISTÊNCIAS")
    build_inconsistencias_table(doc, [
        (1, 'critico', 'AUSÊNCIA DE LAUDO', 'Descrição do problema crítico.'),
        (2, 'alto', 'PROVA FRÁGIL', 'Detalhamento do ponto.'),
    ])

    # Teses
    add_heading(doc, "TESES DEFENSIVAS")
    build_teses_table(doc, [
        (1, 'ATIPICIDADE', 'Fundamentação da tese de atipicidade...'),
        (2, 'DÚVIDA RAZOÁVEL', 'In dubio pro reo...'),
    ])

    # Petição (para peças processuais)
    # build_cabecalho_peca(doc, 'ALEGAÇÕES FINAIS', 'PROC-NUMERO', 'Vara VD', {'reu': 'Nome', 'vítima': 'Nome'})
    # build_requerimento_final(doc, 'Requer-se a absolvição do réu...')

    doc.save(OUTPUT)
    print(f"Salvo: {OUTPUT}")
