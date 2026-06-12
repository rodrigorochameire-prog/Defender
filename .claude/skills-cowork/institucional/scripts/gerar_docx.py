#!/usr/bin/env python3
"""
Gerador de peça DPE-BA (TEMPLATE) — header/footer institucional via python-docx.

Modelo genérico, sem dados de assistido. O corpo abaixo é um esqueleto editável:
preencha o endereçamento, a epígrafe, a qualificação e as seções conforme o caso.

Uso:
    python3 gerar_docx.py [caminho_de_saida.docx]
A logo institucional é lida de ../assets/dpe_logo.png (relativa a este script).
"""

import os
import sys
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt, Twips, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === CONFIGURAÇÃO ===
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_ORIG_PATH = os.path.join(SCRIPT_DIR, "..", "assets", "dpe_logo.png")
LOGO_PATH = os.path.join(tempfile.gettempdir(), "dpe_logo_faded.png")
TEMP_PATH = os.path.join(tempfile.gettempdir(), "_temp_peca.docx")
OUTPUT_PATH = sys.argv[1] if len(sys.argv) > 1 else os.path.join(os.getcwd(), "peca.docx")

# --- Pré-processar logo: aplicar opacidade direto na imagem (60%) ---
from PIL import Image
import numpy as np
img = Image.open(LOGO_ORIG_PATH).convert("RGBA")
arr = np.array(img, dtype=np.float64)
opacity = 0.60
# Misturar com branco na proporção da opacidade
white = np.full_like(arr[:,:,:3], 255.0)
arr[:,:,:3] = arr[:,:,:3] * opacity + white * (1 - opacity)
arr[:,:,3] = 255  # totalmente opaco após blend
result = Image.fromarray(arr.astype(np.uint8)).convert("RGB")
result.save(LOGO_PATH)
print(f"Logo com opacidade {int(opacity*100)}% gerada.")

# === CRIAR DOCUMENTO ===
doc = Document()

# --- Estilos globais ---
style = doc.styles['Normal']
font = style.font
font.name = 'Garamond'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- Configurar seção (margens) ---
section = doc.sections[0]
section.page_width = Twips(11906)
section.page_height = Twips(16838)
section.top_margin = Twips(2552)  # espaço entre logo e corpo levemente reduzido
section.bottom_margin = Twips(1134)
section.left_margin = Twips(1418)   # 2.5cm (mais área útil, ainda adequado para grampeamento)
section.right_margin = Twips(1134)  # 2cm
section.header_distance = Twips(567)
section.footer_distance = Twips(567)

# --- HEADER: logo via python-docx API ---
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
hrun = hp.add_run()
hrun.add_picture(LOGO_PATH, width=Inches(1.777), height=Inches(1.101))

# --- FOOTER: rodapé institucional ---
footer = section.footer
footer.is_linked_to_previous = False

# Limpar parágrafos existentes e criar os do rodapé
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_after = Pt(0)
fp.paragraph_format.space_before = Pt(0)
# Adicionar borda superior ao parágrafo
pPr = fp._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
topBdr = OxmlElement('w:top')
topBdr.set(qn('w:val'), 'single')
topBdr.set(qn('w:sz'), '4')
topBdr.set(qn('w:space'), '1')
topBdr.set(qn('w:color'), '000000')
pBdr.append(topBdr)
pPr.append(pBdr)
# Espaçamento simples
spacing = pPr.find(qn('w:spacing'))
if spacing is None:
    spacing = OxmlElement('w:spacing')
    pPr.append(spacing)
spacing.set(qn('w:line'), '240')
spacing.set(qn('w:lineRule'), 'auto')

frun1 = fp.add_run("Defensoria Pública do Estado da Bahia")
frun1.font.name = 'Arial Narrow'
frun1.font.size = Pt(8)

# Segundo parágrafo do footer
fp2 = footer.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp2.paragraph_format.space_after = Pt(0)
fp2.paragraph_format.space_before = Pt(0)
pPr2 = fp2._p.get_or_add_pPr()
spacing2 = OxmlElement('w:spacing')
spacing2.set(qn('w:line'), '240')
spacing2.set(qn('w:lineRule'), 'auto')
pPr2.append(spacing2)

frun2 = fp2.add_run("7ª Regional da DPE – Camaçari – Bahia.")
frun2.font.name = 'Arial Narrow'
frun2.font.size = Pt(8)

# === HELPERS ===
def add_bold_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(0), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return p

def add_body_paragraph(doc, space_after=Pt(10)):
    """Cria parágrafo do corpo: justificado, recuo 1ª linha, 1.5 espaçamento"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Twips(720)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_run(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return run

def add_title(doc, text):
    """Título de seção: adiciona linha em branco ANTES + título bold"""
    # Linha em branco antes do título (editável pelo usuário)
    add_empty_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Garamond'
    run.font.size = Pt(12)
    return p

def add_empty_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Twips(720)
    return p

# === CONTEÚDO DO DOCUMENTO (ESQUELETO — preencher conforme o caso) ===

# ENDEREÇAMENTO
add_bold_paragraph(doc, "AO JUÍZO DE DIREITO DA [VARA] DA COMARCA DE [COMARCA] – ESTADO DA BAHIA", space_after=Pt(0))

# Duas linhas em branco entre endereçamento e epígrafe
add_empty_line(doc)
add_empty_line(doc)

# EPÍGRAFE
add_bold_paragraph(doc, "Autos nº [NÚMERO DO PROCESSO]", space_after=Pt(20))

# Duas linhas em branco entre epígrafe e qualificação
add_empty_line(doc)
add_empty_line(doc)

# QUALIFICAÇÃO + TIPO DA PEÇA (inline)
p = add_body_paragraph(doc)
add_run(p, "[NOME DO ASSISTIDO]", bold=True)
add_run(p, ", [qualificação: nacionalidade, estado civil, profissão, RG, CPF, já qualificado nos autos], representado pela Defensoria Pública do Estado da Bahia, com fundamento no art. 134 da Constituição da República, por meio do defensor público subscritor, vem respeitosamente perante Vossa Excelência apresentar o presente ")
add_run(p, "[TIPO DA PEÇA]", bold=True)
add_run(p, " (fundamento legal), com fundamento nos fatos e razões de direito a seguir expostos.")

# I – DOS FATOS
add_title(doc, "I – DOS FATOS")
p = add_body_paragraph(doc)
add_run(p, "[Descrever sucintamente os fatos relevantes do caso.]")

# II – DO DIREITO
add_title(doc, "II – DO DIREITO")
p = add_body_paragraph(doc)
add_run(p, "[Desenvolver a fundamentação jurídica da tese defensiva, com citação de dispositivos legais e precedentes.]")

# III – DOS PEDIDOS
add_title(doc, "III – DOS PEDIDOS")
p = add_body_paragraph(doc)
add_run(p, "[Formular os pedidos, em ordem principal e subsidiária.]")

# FECHO
p = add_body_paragraph(doc)
add_run(p, "Nesses termos, pede deferimento.")

# DATA - justificado com recuo (como no corpo, mesmo espaçamento)
p = add_body_paragraph(doc)
add_run(p, "Camaçari – BA, [DATA].")

# Linha em branco
doc.add_paragraph()

# ASSINATURA - centralizada
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing_rule = 1  # single
add_run(p, "Rodrigo Rocha Meire", bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing_rule = 1
add_run(p, "Defensor Público", bold=True)


# === SALVAR TEMP E PÓS-PROCESSAR ===
doc.save(TEMP_PATH)
print("Documento base criado com header/footer via python-docx.")

# Copiar diretamente (opacidade já está na imagem, sem depender de alphaModFix)
import shutil
shutil.move(TEMP_PATH, OUTPUT_PATH)
print(f"OK: {OUTPUT_PATH}")
