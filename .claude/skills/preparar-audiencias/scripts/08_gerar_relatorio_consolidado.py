#!/usr/bin/env python3
"""
Gera o relatório consolidado de audiências de um dia em PDF (com .docx editável).
Lê /tmp/pauta-<YYYY-MM-DD>.json e produz Pauta de Audiências em
~/.../Atendimentos/Pauta de Audiências - <DD MES YYYY>.{docx,pdf}.

Uso:
  python3 .claude/skills-cowork/preparar-audiencias/scripts/08_gerar_relatorio_consolidado.py 2026-05-05
"""
import json
import subprocess
import sys
from datetime import datetime
from pathlib import Path

import numpy as np
from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


LOGO_SRC = "/Users/rodrigorochameire/Projetos/Defender/.claude/skills-cowork/vvd/assets/dpe_logo.png"
LOGO_FADED = "/tmp/dpe_logo_faded.png"
OUT_DIR = Path("/Users/rodrigorochameire/Library/CloudStorage/GoogleDrive-rodrigorochameire@gmail.com/Meu Drive/1 - Defensoria 9ª DP/5 - Operacional/Atendimentos")
SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

PALETA = {
    "VVD_CAMACARI": {"primaria": "A16207", "secundaria": "78716C"},
    "JURI_CAMACARI": {"primaria": "047857", "secundaria": "57534E"},
    "GRUPO_JURI": {"primaria": "047857", "secundaria": "57534E"},
    "EXECUCAO_PENAL": {"primaria": "1E40AF", "secundaria": "57534E"},
    "SUBSTITUICAO": {"primaria": "475569", "secundaria": "78716C"},
    "default": {"primaria": "A16207", "secundaria": "78716C"},
}

NOMES_MES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]


def fade_logo(src: str, dst: str, opacity: float = 0.60) -> None:
    img = Image.open(src).convert("RGBA")
    arr = np.array(img, dtype=np.float64)
    white = np.full_like(arr[:, :, :3], 255.0)
    arr[:, :, :3] = arr[:, :, :3] * opacity + white * (1 - opacity)
    arr[:, :, 3] = 255
    Image.fromarray(arr.astype(np.uint8)).convert("RGB").save(dst)


def fmt_hour(iso, horario):
    if horario:
        return horario
    if not iso:
        return "—"
    try:
        return datetime.fromisoformat(iso.replace("Z", "+00:00")).strftime("%H:%M")
    except Exception:
        return "—"


def add_para(doc_or_cell, text, *, bold=False, size=None, align=None,
              space_after=None, space_before=None, indent_first=None,
              color=None, italic=False, line_spacing=1.5):
    p = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, "add_paragraph") else doc_or_cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    if space_after is not None: pf.space_after = Pt(space_after)
    if space_before is not None: pf.space_before = Pt(space_before)
    if indent_first is not None: pf.first_line_indent = Twips(indent_first)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size: run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_section_title(doc, title, color="A16207"):
    add_para(doc, "", space_after=0)
    return add_para(doc, title, bold=True, size=14, color=color,
                    space_before=18, space_after=12,
                    align=WD_ALIGN_PARAGRAPH.LEFT)


# Painel de depoentes — render
ROTULOS_INTIMACAO = {
    "intimado": ("INTIMADO", "059669"),
    "nao_intimado": ("NÃO INTIMADO", "DC2626"),
    "pendente": ("PENDENTE", "EA580C"),
    "dispensada": ("DISPENSADA", "78716C"),
    "desconhecido": ("DESCONHECIDO", "78716C"),
}
ROTULOS_COMPARECIMENTO = {
    "compareceu": ("COMPARECEU", "059669"),
    "nao_compareceu": ("NÃO COMPARECEU", "DC2626"),
    "nao_verificado": ("A VERIFICAR", "EA580C"),
    "dispensada": ("DISPENSADA", "78716C"),
    "ouvido_anteriormente": ("OUVIDO ANT.", "0284C7"),
    "substituida": ("SUBSTITUÍDA", "78716C"),
    "contraditada": ("CONTRADITADA", "78716C"),
}
ROTULOS_TIPO = {
    "ofendida": "Ofendida",
    "testemunha_acusacao": "Test. Acus.",
    "testemunha_defesa": "Test. Defesa",
    "informante": "Informante",
    "interrogando": "Interrogando",
    "vitima_indireta": "Vít. Indireta",
    "perito": "Perito",
    "assistente_tecnico": "Assist. Téc.",
}
ROTULOS_MOTIVO = {
    "nao_localizado": "não localizado",
    "mandado_nao_cumprido": "mandado não cumprido",
    "endereco_invalido": "endereço inválido",
    "em_diligencia": "em diligência",
    "recusa_recebimento": "recusa",
    "precatoria_devolvida": "precatória devolvida",
    "precatoria_pendente": "precatória pendente",
    "mandado_nao_emitido": "mandado não emitido",
    "falta_de_informacoes": "falta de informações",
}


def render_painel_depoentes(doc, depoentes: list[dict], subtipo: str):
    if not depoentes:
        if subtipo in ("custodia", "qualificacao"):
            add_para(doc, f"Não se aplica — audiência de {subtipo} (sem depoentes arrolados).",
                      italic=True, size=10, color="78716C", space_after=8)
            return
        else:
            add_para(doc, "⚠ Painel de depoentes não preenchido. Pendência: levantar arrolados nos autos.",
                      italic=True, size=10, color="DC2626", space_after=8)
            return

    add_para(doc, "Painel de Depoentes", bold=True, size=11, color="78716C",
              space_before=6, space_after=4)

    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Light Grid"
    hdr = tbl.rows[0].cells
    headers = ["#", "Nome", "Tipo", "Intimação", "Comparecimento", "Já ouvido", "Forma", "Observação"]
    for c, label in zip(hdr, headers):
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(label); r.bold = True; r.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)

    for i, d in enumerate(depoentes, 1):
        row = tbl.add_row().cells
        row[0].text = str(i)

        row[1].text = d.get("nome", "—")

        row[2].text = ROTULOS_TIPO.get(d.get("tipo"), d.get("tipo", "—"))

        intim_label, intim_color = ROTULOS_INTIMACAO.get(d.get("intimacao"), ("—", "78716C"))
        cell = row[3]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(intim_label); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(intim_color)
        if d.get("motivo_nao_intimacao"):
            p.add_run(f"\n{ROTULOS_MOTIVO.get(d['motivo_nao_intimacao'], d['motivo_nao_intimacao'])}").font.size = Pt(8)

        comp_label, comp_color = ROTULOS_COMPARECIMENTO.get(d.get("comparecimento"), ("—", "78716C"))
        cell = row[4]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(comp_label); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(comp_color)

        ja = d.get("ja_ouvido")
        if ja and ja.get("sim"):
            txt = f"SIM\n{ja.get('data', '?')}\n{ja.get('peca', '')}"
            row[5].text = txt
        else:
            row[5].text = "Não"

        row[6].text = (d.get("forma") or "—").replace("_", " ")
        row[7].text = (d.get("observacao") or "—")[:200]

        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    if not r.font.size:
                        r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(0)


def gerar(dia: str):
    pauta = json.loads(Path(f"/tmp/pauta-{dia}.json").read_text())
    auds = pauta["audiencias"]

    fade_logo(LOGO_SRC, LOGO_FADED)

    # Detectar atribuição predominante
    cnt = {}
    for a in auds:
        cnt[a.get("atribuicao", "default")] = cnt.get(a.get("atribuicao", "default"), 0) + 1
    atribuicao = max(cnt, key=cnt.get) if cnt else "default"
    cor_primaria = PALETA.get(atribuicao, PALETA["default"])["primaria"]
    cor_secundaria = PALETA.get(atribuicao, PALETA["default"])["secundaria"]

    doc = Document()

    # Margens
    for s in doc.sections:
        s.top_margin = Twips(2552)
        s.bottom_margin = Twips(1134)
        s.left_margin = Twips(1418)
        s.right_margin = Twips(1134)
        s.header_distance = Twips(567)
        s.footer_distance = Twips(567)

    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.font.size = Pt(12)

    # Header — logo
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run().add_picture(LOGO_FADED, width=Inches(1.777), height=Inches(1.101))

    # Footer
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        for k, v in [("w:val", "single"), ("w:sz", "4"), ("w:space", "1"), ("w:color", "000000")]:
            top.set(qn(k), v)
        pBdr.append(top)
        pPr.append(pBdr)
        break
    f1 = footer.paragraphs[0]; f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f1.add_run("Defensoria Pública do Estado da Bahia"); fr.font.name = "Arial Narrow"; fr.font.size = Pt(8)
    f2 = footer.add_paragraph(); f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = f2.add_run("7ª Regional da DPE – Camaçari – Bahia."); fr2.font.name = "Arial Narrow"; fr2.font.size = Pt(8)

    # Cabeçalho
    dt = datetime.fromisoformat(dia)
    nome_mes = NOMES_MES[dt.month - 1].upper()
    titulo = add_para(doc, f"PAUTA DE AUDIÊNCIAS — {dt.day:02d} DE {nome_mes} DE {dt.year}",
                      bold=True, size=14, color=cor_primaria,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
    for r in titulo.runs: r.font.underline = True
    add_para(doc, "Defensoria Pública · 7ª Regional · Camaçari/BA",
              italic=True, size=11, color=cor_secundaria,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # KPIs
    total = len(auds)
    agendadas = sum(1 for a in auds if a.get("status") == "agendada")
    concluidas = sum(1 for a in auds if a.get("status") == "concluido")
    canceladas = sum(1 for a in auds if a.get("status") == "cancelada")
    com_painel = sum(1 for a in auds if (a.get("registro_audiencia") or {}).get("depoentes"))
    com_resumo = sum(1 for a in auds if a.get("resumo_defesa"))

    kpi = doc.add_table(rows=1, cols=6)
    kpi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells = kpi.rows[0].cells
    items = [
        (str(total), "Audiências"),
        (str(agendadas), "Agendadas"),
        (str(concluidas), "Concluídas"),
        (str(canceladas), "Canceladas"),
        (f"{com_painel}/{total}", "Painel"),
        (f"{com_resumo}/{total}", "Resumo"),
    ]
    for cell, (n, lbl) in zip(cells, items):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = cell.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(n); r1.bold = True; r1.font.size = Pt(20)
        r1.font.color.rgb = RGBColor.from_string(cor_primaria)
        p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(lbl); r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(cor_secundaria)
    add_para(doc, "", space_after=12)

    # Sinóptico
    add_section_title(doc, "I. Sinóptico do Dia", color=cor_primaria)
    tbl = doc.add_table(rows=1, cols=6); tbl.style = "Light Grid"
    hdr = tbl.rows[0].cells
    for c, lbl in zip(hdr, ["Hora", "Tipo", "Assistido", "Autos", "Status", "Painel"]):
        c.text = ""; p = c.paragraphs[0]
        r = p.add_run(lbl); r.bold = True; r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for a in auds:
        row = tbl.add_row().cells
        row[0].text = fmt_hour(a.get("data_audiencia"), a.get("horario"))
        row[1].text = a.get("tipo", "—")
        row[2].text = a.get("assistido_nome", "—")
        row[3].text = a.get("numero_autos", "—") or "—"
        row[4].text = a.get("status", "—")
        tem_painel = bool((a.get("registro_audiencia") or {}).get("depoentes"))
        row[5].text = "✓" if tem_painel else "⚠"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(0)
    add_para(doc, "", space_after=12)

    # Detalhamento
    add_section_title(doc, "II. Detalhamento por Audiência", color=cor_primaria)
    for idx, a in enumerate(auds, 1):
        hora = fmt_hour(a.get("data_audiencia"), a.get("horario"))
        nome = a.get("assistido_nome", "—")
        tipo = a.get("tipo", "—")
        add_para(doc, f"#{idx} · {hora} · {tipo} — {nome}",
                  bold=True, size=12, color=cor_primaria,
                  space_before=12, space_after=4)

        status = (a.get("status") or "").upper() or "—"
        status_color = {"AGENDADA": "059669", "CONCLUIDO": "0284C7", "CANCELADA": "DC2626"}.get(status, "78716C")
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.2; p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run("Status: "); r1.font.size = Pt(10); r1.font.color.rgb = RGBColor.from_string("78716C")
        r2 = p.add_run(status); r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor.from_string(status_color)

        # Dados estruturais
        rg = a.get("registro_audiencia") or {}
        linhas = [
            ("Autos", a.get("numero_autos") or "—"),
            ("Classe", a.get("classe_processual") or "—"),
            ("Vara/Comarca", f"{a.get('vara') or '—'} · {a.get('comarca') or '—'}"),
            ("Local/Sala", f"{a.get('local') or '—'} · Sala {a.get('sala') or '—'}"),
            ("Imputação", (rg.get("imputacao") or {}).get("principal") or "—"),
        ]
        for label, val in linhas:
            p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(0)
            r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor.from_string("57534E")
            r2 = p.add_run(str(val)); r2.font.size = Pt(10)

        # PAINEL DE DEPOENTES — sempre presente
        render_painel_depoentes(doc, rg.get("depoentes", []), rg.get("subtipo_audiencia", "indefinido"))

        # Resumo de defesa
        if a.get("resumo_defesa"):
            add_para(doc, "Resumo de defesa:", bold=True, size=11, color=cor_secundaria,
                      space_before=8, space_after=2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.4; p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Twips(360)
            r = p.add_run(a["resumo_defesa"]); r.font.size = Pt(11)

        # Tese principal
        if rg.get("tese_defesa", {}).get("principal"):
            add_para(doc, "Tese principal:", bold=True, size=10, color=cor_secundaria,
                      space_before=6, space_after=2)
            add_para(doc, rg["tese_defesa"]["principal"], italic=True, size=11,
                      space_after=6, line_spacing=1.3)

        # Pontos críticos
        if rg.get("pontos_criticos"):
            add_para(doc, "Pontos críticos:", bold=True, size=10, color=cor_secundaria,
                      space_before=6, space_after=2)
            for pt in rg["pontos_criticos"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.3
                r = p.add_run(pt); r.font.size = Pt(11)

        # Perguntas estratégicas
        if rg.get("perguntas_estrategicas"):
            add_para(doc, "Perguntas estratégicas:", bold=True, size=10,
                      color=cor_secundaria, space_before=6, space_after=2)
            for grupo, perguntas in rg["perguntas_estrategicas"].items():
                if not perguntas: continue
                add_para(doc, f"Para {grupo.replace('_', ' ')}:", bold=True, size=10,
                          space_after=0, line_spacing=1.2)
                for q in perguntas:
                    p = doc.add_paragraph(style="List Number")
                    p.paragraph_format.line_spacing = 1.3
                    r = p.add_run(q); r.font.size = Pt(11)

        # Orientação ao defendido
        if rg.get("orientacao_assistido"):
            add_para(doc, "Orientação ao defendido:", bold=True, size=10,
                      color=cor_secundaria, space_before=6, space_after=2)
            add_para(doc, rg["orientacao_assistido"], size=11, space_after=8,
                      line_spacing=1.4)

        # Documentos relevantes
        if rg.get("documentos_relevantes"):
            add_para(doc, "Documentos relevantes:", bold=True, size=10,
                      color=cor_secundaria, space_before=6, space_after=2)
            for d in rg["documentos_relevantes"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.2
                r = p.add_run(f"{d.get('data', '—')} · {d.get('tipo', '')} · ID {d.get('id_pje', '?')} (Pág. {d.get('fl', '?')}) · {d.get('descricao', '')}")
                r.font.size = Pt(10)

        add_para(doc, "", space_after=8)

    # Pendências
    pend_geral = []
    for a in auds:
        rg = a.get("registro_audiencia") or {}
        if not rg.get("depoentes") and a.get("status") != "cancelada":
            pend_geral.append(f"#{a['id']} {a.get('assistido_nome')} — painel de depoentes vazio")
        if not a.get("resumo_defesa") and a.get("status") != "cancelada":
            pend_geral.append(f"#{a['id']} {a.get('assistido_nome')} — resumo de defesa vazio")
        for p in (rg.get("pendencias") or []):
            pend_geral.append(f"#{a['id']} {a.get('assistido_nome')} — {p}")

    if pend_geral:
        add_section_title(doc, "III. Pendências do Dia", color="DC2626")
        for p_txt in pend_geral:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = 1.3
            r = p.add_run(p_txt); r.font.size = Pt(11)

    # Footer textual
    gerado = datetime.now().strftime("%d/%m/%Y às %H:%M")
    add_para(doc, "", space_after=12)
    add_para(doc, f"Pauta gerada em {gerado} a partir do OMBUDS · {len(auds)} audiências.",
              italic=True, size=9, color="A8A29E",
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Salvar
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    nome_pdf = f"Pauta de Audiências - {dt.day:02d} {NOMES_MES[dt.month-1]} {dt.year}"
    out_docx = OUT_DIR / f"{nome_pdf}.docx"
    out_pdf = OUT_DIR / f"{nome_pdf}.pdf"
    doc.save(out_docx)
    print(f"DOCX: {out_docx}")

    res = subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf",
                          "--outdir", str(OUT_DIR), str(out_docx)],
                         capture_output=True, text=True, timeout=180)
    if res.returncode != 0:
        print("LIBRE_ERR:", res.stderr)
    else:
        print(f"PDF: {out_pdf}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    gerar(sys.argv[1])
