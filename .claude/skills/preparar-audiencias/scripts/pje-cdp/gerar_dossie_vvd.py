#!/usr/bin/env python3
"""
Gera o DOSSIÊ ESTRATÉGICO VVD (.docx, paleta amber, Padrão Defender v2)
a partir do JSON estruturado produzido pela análise.

Uso: gerar_dossie_vvd.py dossie.json saida.docx
Conversão a PDF: soffice --headless --convert-to pdf saida.docx
"""
import json, sys, os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Twips, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paleta VVD (amber) — Padrão Defender v2
DARK = "6B4D2B"; ACCENT = "C8A84E"; SUBTLE = "FAF8F2"; SOFT = "F5E6A3"
TITLE = "5C3D1A"; MUTED = "8B7355"; BORDER = "E5E5E5"
RED_BG = "FEE2E2"; RED = "DC2626"; GREEN_BG = "D1FAE5"; YELLOW_BG = "FEF9C3"

LOGO = "/Users/rodrigorochameire/Meu Drive/1 - Defensoria 9ª DP/11 - Arquivo & sistema/Skills - harmonizacao/vvd/assets/dpe_logo.png"


def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, side, color, sz='8'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders'); tcPr.append(borders)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
    borders.append(el)


def all_borders(table, color=BORDER):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        borders.append(el)
    tblPr.append(borders)


def cell_text(cell, text, bold=False, size=9, color=None, align=None, font='Verdana'):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font; run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return p


def heading(doc, text):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    set_cell_bg(c, "FFFFFF"); set_cell_border(c, 'left', ACCENT, '24')
    cell_text(c, text, bold=True, size=12, color=TITLE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def subheading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = 'Verdana'; r.font.size = Pt(10.5)
    r.font.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)


def para(doc, text, size=10, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(str(text)); r.font.name = 'Verdana'; r.font.size = Pt(size); r.italic = italic
    return p


def bullet(doc, text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    if label:
        r = p.add_run(label + ": "); r.font.bold = True
        r.font.name = 'Verdana'; r.font.size = Pt(10)
    r = p.add_run(str(text)); r.font.name = 'Verdana'; r.font.size = Pt(10)


def kv_table(doc, pairs, col1=2400, col2=6800):
    t = doc.add_table(rows=0, cols=2)
    all_borders(t)
    for k, v in pairs:
        row = t.add_row()
        set_cell_bg(row.cells[0], SUBTLE)
        cell_text(row.cells[0], k, bold=True, size=9, color=TITLE)
        cell_text(row.cells[1], v if v not in (None, "") else "—", size=9)
    return t


def main():
    data = json.loads(Path(sys.argv[1]).read_text())
    out = Path(sys.argv[2])

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Verdana'; style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sec = doc.sections[0]
    sec.top_margin = Twips(1418); sec.bottom_margin = Twips(1134)
    sec.left_margin = Twips(1418); sec.right_margin = Twips(1134)

    # Header com logo
    if os.path.exists(LOGO):
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.add_run().add_picture(LOGO, width=Inches(1.4))
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Defensoria Pública do Estado da Bahia — 7ª Regional da DPE – Camaçari – Bahia.")
    fr.font.name = 'Arial Narrow'; fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)

    # Capa
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DOSSIÊ ESTRATÉGICO DE DEFESA"); r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TITLE); r.font.name = 'Verdana'
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("assistido", "")); r.font.size = Pt(14); r.font.name = 'Verdana'
    aud = data.get("audiencia", {})
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{aud.get("tipo","")} — 11/06/2026 às {aud.get("horario","")} · Vara de Violência Doméstica de Camaçari')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(MUTED); r.font.name = 'Verdana'

    # KPIs (4 cards)
    kpis = data.get("kpis", [])[:4]
    if kpis:
        t = doc.add_table(rows=2, cols=len(kpis)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        all_borders(t)
        for i, k in enumerate(kpis):
            set_cell_bg(t.rows[0].cells[i], SUBTLE)
            cell_text(t.rows[0].cells[i], str(k.get("valor", "")), bold=True, size=20,
                      color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_text(t.rows[1].cells[i], k.get("label", ""), size=8, color=MUTED,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # PARTE I — resumo executivo
    heading(doc, "I. RESUMO EXECUTIVO ESTRATÉGICO")
    for px in data.get("resumo_executivo", []):
        para(doc, px)

    # Painel de controle
    heading(doc, "II. PAINEL DE CONTROLE DO CASO")
    kv_table(doc, [(k, v) for k, v in data.get("painel_controle", {}).items()])
    doc.add_paragraph()

    # Imputação / medidas
    imp = data.get("imputacao", {})
    if imp:
        subheading(doc, "Imputação / objeto")
        kv_table(doc, [(k.replace("_", " ").capitalize(),
                        ", ".join(v) if isinstance(v, list) else v)
                       for k, v in imp.items() if v])
        doc.add_paragraph()
    if data.get("medidas_mpu"):
        subheading(doc, "Medidas protetivas")
        for m in data["medidas_mpu"]:
            bullet(doc, m)

    # PARTE III — Painel de depoentes (REGRA DE OURO)
    heading(doc, "III. PAINEL DE DEPOENTES")
    deps = data.get("depoentes", [])
    if deps:
        cols = ["Nome", "Tipo", "Intimação", "Motivo", "Compareci/o", "Já ouvido", "Forma", "Obs"]
        t = doc.add_table(rows=1, cols=len(cols)); all_borders(t)
        for i, c in enumerate(cols):
            set_cell_bg(t.rows[0].cells[i], DARK)
            cell_text(t.rows[0].cells[i], c, bold=True, size=8, color="FFFFFF")
        for d in deps:
            row = t.add_row()
            jo = d.get("ja_ouvido")
            jo_txt = f'SIM ({jo.get("data","")} {jo.get("peca","")})' if jo and jo.get("sim") else "NÃO"
            vals = [d.get("nome", ""), d.get("tipo", "").replace("_", " "),
                    d.get("intimacao", ""), d.get("motivo_nao_intimacao") or "—",
                    d.get("comparecimento", ""), jo_txt, d.get("forma", ""),
                    d.get("observacao", "")]
            for i, v in enumerate(vals):
                cell_text(row.cells[i], v, size=7.5)
            intim = (d.get("intimacao") or "").lower()
            bg = GREEN_BG if intim == "intimado" else (RED_BG if intim == "nao_intimado" else YELLOW_BG if intim == "pendente" else "FFFFFF")
            set_cell_bg(row.cells[2], bg)
    else:
        para(doc, "Não se aplica — sem depoentes arrolados para este ato.", italic=True)
    doc.add_paragraph()

    # Cronologia
    if data.get("cronologia"):
        heading(doc, "IV. CRONOLOGIA PROCESSUAL")
        t = doc.add_table(rows=0, cols=2); all_borders(t)
        for ev in data["cronologia"]:
            row = t.add_row()
            set_cell_bg(row.cells[0], SUBTLE)
            cell_text(row.cells[0], f'{ev.get("marcador","•")} {ev.get("data","")}', bold=True, size=8.5, color=TITLE)
            cell_text(row.cells[1], ev.get("evento", ""), size=8.5)
        doc.add_paragraph()

    # Fragilidades
    if data.get("pontos_criticos"):
        heading(doc, "V. PONTOS CRÍTICOS / FRAGILIDADES")
        for i, pc in enumerate(data["pontos_criticos"], 1):
            bullet(doc, pc, label=f"{i}")

    # Teses
    if data.get("teses"):
        heading(doc, "VI. TESES DEFENSIVAS")
        for tese in data["teses"]:
            t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]
            set_cell_border(c, 'left', ACCENT, '20'); set_cell_bg(c, SUBTLE)
            viab = tese.get("viabilidade", "")
            blocks = {"ALTA": "■■■■□", "MÉDIA": "■■■□□", "MEDIA": "■■■□□", "BAIXA": "■■□□□"}.get(viab.upper(), "")
            cell_text(c, f'{tese.get("nome","")}   {blocks} {viab}', bold=True, size=10, color=TITLE)
            for key, lab in (("fundamento", "Fundamento"), ("elementos", "Elementos"), ("riscos", "Riscos")):
                v = tese.get(key)
                if v:
                    bullet(doc, "; ".join(v) if isinstance(v, list) else v, label=lab)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        if data.get("narrativa"):
            subheading(doc, "Narrativa defensiva")
            para(doc, data["narrativa"], italic=True)

    # Perguntas
    pe = data.get("perguntas_estrategicas", {})
    if any(pe.values() if isinstance(pe, dict) else []):
        heading(doc, "VII. PERGUNTAS ESTRATÉGICAS")
        for alvo, qs in pe.items():
            if not qs: continue
            subheading(doc, alvo.replace("_", " ").capitalize())
            for q in qs:
                bullet(doc, q)

    # Orientação
    if data.get("orientacao_assistido"):
        heading(doc, "VIII. ORIENTAÇÃO AO ASSISTIDO")
        para(doc, data["orientacao_assistido"])

    # Requerimentos orais
    if data.get("requerimentos_orais"):
        heading(doc, "IX. REQUERIMENTOS ORAIS PRONTOS")
        for rq in data["requerimentos_orais"]:
            t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]
            set_cell_bg(c, SUBTLE); set_cell_border(c, 'left', DARK, '16')
            cell_text(c, rq, size=9)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Cenários
    cen = data.get("cenarios", {})
    if cen:
        heading(doc, "X. CENÁRIOS")
        for nome, txt in cen.items():
            bullet(doc, txt, label=nome.replace("_", " ").capitalize())

    # Providências
    prov = data.get("providencias", {})
    if prov:
        heading(doc, "XI. PROVIDÊNCIAS")
        for fase, items in prov.items():
            if not items: continue
            subheading(doc, fase.replace("_", " ").capitalize())
            for it in items:
                bullet(doc, f"☐ {it}")

    # Pendências
    if data.get("pendencias"):
        heading(doc, "XII. PENDÊNCIAS / LACUNAS")
        for pd_ in data["pendencias"]:
            bullet(doc, pd_)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"OK {out}")


if __name__ == "__main__":
    main()
